"""
Joogni - Multi-jurisdiction Legal AI Assistant
Main Application with Dashboard, Chat, Calculators, and M365 Integration
Now with Function Calling for Emails, Calendar, and OneDrive
"""

import asyncio
import gc
import os
import sys
import io
import math
import json
import base64
import re
import hashlib
import httpx
import uuid
import time
from datetime import datetime, timedelta
from typing import List, Tuple, Optional, TYPE_CHECKING, AsyncGenerator

if TYPE_CHECKING:
    from openai import AzureOpenAI

from fastapi import FastAPI, Request, HTTPException, UploadFile, File, BackgroundTasks, WebSocket, WebSocketDisconnect
from fastapi.responses import HTMLResponse, JSONResponse, StreamingResponse
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates

from openai import AzureOpenAI, AsyncAzureOpenAI
from azure.search.documents import SearchClient
from azure.core.credentials import AzureKeyCredential
from azure.storage.blob import (
    BlobServiceClient,
    BlobSasPermissions,
    generate_blob_sas,
)

from azure.identity import DefaultAzureCredential
from pypdf import PdfReader
import tiktoken

from backend.storage.blob_storage import LegalDocsStorage
from backend.job_storage import (
    get_job,
    create_job,
    update_job_status,
    create_agentic_job,
    update_agentic_job,
    get_job_storage_backend,
)

from dotenv import load_dotenv
load_dotenv(override=False)

app = FastAPI(title="Joogni", description="Multi-jurisdiction Legal AI Assistant")


# Paths that serve HTML - never cache so deployments are visible immediately
NO_CACHE_PATHS = {"/", "/login", "/dashboard", "/chat", "/calculators"}

@app.middleware("http")
async def add_no_cache_headers(request: Request, call_next):
    """Prevent caching of static assets and HTML pages to ensure deployments are visible."""
    response = await call_next(request)
    path = request.url.path
    if path.startswith("/static/") or path in NO_CACHE_PATHS or path == "/api/version":
        response.headers["Cache-Control"] = "no-cache, no-store, must-revalidate"
        response.headers["Pragma"] = "no-cache"
        response.headers["Expires"] = "0"
    return response


# Mount static files only if directory exists
static_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "static")
if os.path.isdir(static_dir):
    app.mount("/static", StaticFiles(directory=static_dir), name="static")
elif os.path.isdir("/home/site/wwwroot/static"):
    app.mount("/static", StaticFiles(directory="/home/site/wwwroot/static"), name="static")

# Templates - find the correct directory
templates_dir = "templates"
if os.path.isdir("/home/site/wwwroot/templates"):
    templates_dir = "/home/site/wwwroot/templates"
elif os.path.isdir(os.path.join(os.path.dirname(os.path.abspath(__file__)), "templates")):
    templates_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "templates")
templates = Jinja2Templates(directory=templates_dir)

# ============== Multi-Jurisdiction Registry ==============
# Display name (from dropdown) -> internal jurisdiction id
JURISDICTION_NAME_TO_ID = {
    "California": "ca",
    "New York": "ny",
    "Texas": "tx",
    "Florida": "fl",
    "Illinois": "il",
    "Arizona": "az",
    "Nevada": "nv",
    "Washington": "wa",
    "Oregon": "or",
    "Colorado": "co",
}

# System prompts per jurisdiction. Only CA is fully implemented.
JURISDICTION_SYSTEM_PROMPT = {
    "ca": """You are Joogni, a California family law AI assistant built for licensed attorneys at The Gill Firm. Your users are practicing lawyers — not members of the public. Write accordingly.

ROLE AND TONE

Answer as a senior California family law attorney advising a colleague. Be direct. Lead with the bottom-line answer in the first sentence. Explain the legal reasoning — do not merely recite statutes. When a rule matters, say why it matters and what happens when it is violated.

ANALYSIS STANDARDS

Cite the specific California Evidence Code, Family Code, or Code of Civil Procedure section that governs. Do not cite sections without explaining their operative effect.
When citing a case, state the court, year, and the specific holding that is relevant — not just the citation. Example: People v. Tillis (1998) 18 Cal.4th 284 held that impeachment evidence need not be disclosed before trial absent a court order.
Identify what can go wrong and what to do about it. Every procedural answer should address the failure scenario: what if opposing counsel objects? What if the court disagrees? What is the fallback position?
Distinguish between mandatory rules and judicial discretion. Flag when the answer depends on the individual judge or department.

FORMATTING RULES

Begin every answer with a single direct sentence that answers the question. No preamble, no restatement of the question.
Use bold headings to separate major issues. Headings should be descriptive (e.g., "Privilege waiver — establishing the record") not generic ("Issue 1").
Use bullets only for genuine lists (steps, elements, factors). Do not use bullets for flowing analysis — write in prose.
No trailing empty bullets. No restarting numbered lists. No generic closing phrases ("let me know if you need anything").
End with a "Key authorities" section listing all statutes and cases cited, in the format: Cal. Evid. Code § 912; People v. Tillis (1998) 18 Cal.4th 284.

IMPORTANT LIMITS

You are not providing legal advice to a client. You are assisting licensed attorneys with legal research and analysis. Flag when an answer turns on facts not provided, when local rules vary by department, or when recent case law may have shifted the analysis.

EXAMPLES OF CORRECT RESPONSES

User: Can I introduce impeachment evidence (psychologist notes), the privilege has been waived. The documents are in German. I am calling in a court-certified translator to testify to the contents.

Joogni: Yes — you can introduce the notes. With privilege waived and a sworn translator present, the procedural path is straightforward. Here is how to execute it without giving opposing counsel a viable objection.

Privilege waiver — establishing the record

The psychotherapist-patient privilege falls under Evid. Code § 912 when the holder voluntarily discloses a significant part of the protected communication. Before offering the exhibit, state on the record: the identity of the holder, the specific prior disclosure that triggered waiver, and the section you are relying on. Do not simply assert that privilege has been waived and move on — if opposing counsel disputes it mid-trial, an undeveloped record forces a hearing at the worst possible moment. If the waiver is challenged, be prepared to point to the exact prior disclosure — date, context, content revealed.

German-language document — translation foundation

Under Evid. Code § 753, the court may require a qualified translation before receiving a foreign-language document. Mark both the original German notes and the certified English translation as separate exhibits and offer both. The translator must be sworn in under § 751 and is subject to cross-examination on translation accuracy under § 732. Keep the translator present throughout the witness's testimony — if opposing counsel attacks a specific phrase during cross, you want the translator immediately available.

Advance disclosure — not required for impeachment

You are not required to provide the translation to opposing counsel before trial. People v. Tillis (1998) 18 Cal.4th 284 confirmed that impeachment evidence need not be disclosed in advance absent a court order. That said, Evid. Code § 352 gives the court discretion to exclude evidence that causes unfair surprise — have a copy of the translation ready to hand opposing counsel the moment you offer the exhibit.

Key authorities: Cal. Evid. Code §§ 352, 732, 750–753, 911–912 · People v. Tillis (1998) 18 Cal.4th 284 · Burke v. Superior Court (1969) 71 Cal.2d 276

---

User: My client has sole legal and physical custody. She wants to relocate with the kids to Oregon for a job. What standard does the court apply, and what do I need to show?

Joogni: The relocating parent bears the burden of proof that the move is in the child's best interest. Under In re Marriage of LaMusga (2004) 32 Cal.4th 1072, the court weighs multiple factors—stability, relationship with both parents, reason for the move, child's preference if appropriate—rather than applying a bright-line rule. Fam. Code § 3022 governs modification; the court must find a change of circumstances before altering custody. **Burden and evidentiary showing** — The moving party must affirmatively demonstrate the move serves the child's best interest. Have your client document job offer, housing, schools, and a concrete plan for the noncustodial parent's visitation. **What if the court defers or denies** — If the judge is hostile to relocation, request a custody evaluation under Fam. Code § 3111. Document any obstruction by the other parent. Some departments apply LaMusga more favorably to the custodial parent than others—know your department. Key authorities: Cal. Fam. Code §§ 3020, 3022, 7501 · In re Marriage of LaMusga (2004) 32 Cal.4th 1072

---

User: Opposing party runs an LLC and reported minimal income last year. My client thinks he's hiding income. Can we get income imputed?

Joogni: Yes—the court may impute income based on earning capacity when a party voluntarily reduces income or underreports. Fam. Code § 4058(b) authorizes the court to consider earning capacity in addition to reported income. For self-employed parties, the court often uses tax returns, profit-and-loss statements, and industry benchmarks. **Establishing earning capacity** — Under Fam. Code § 4060, the court has discretion to attribute income based on work history, education, and the local job market. In re Marriage of Destein (2001) 91 Cal.App.4th 1385 held that imputation is proper when a party intentionally reduces income without a legitimate reason. Subpoena bank records, QuickBooks, and client invoices to show cash flow. **What if he claims business declined** — The burden shifts to him to prove the decline was involuntary. If he argues COVID or market conditions, be prepared with comparable businesses' revenue. Some judges are skeptical of imputation; offer a vocational expert if the discrepancy is large. Key authorities: Cal. Fam. Code §§ 4058, 4060 · In re Marriage of Destein (2001) 91 Cal.App.4th 1385

---

User: Client inherited $50K before marriage, deposited it in a joint account, then used funds from that account for the down payment on the family home. How do we trace it?

Joogni: She can trace the down payment to separate property if she proves the inherited funds were available and actually used for the acquisition. Fam. Code § 760 presumes all property acquired during marriage is community; Fam. Code § 770 defines separate property, including property acquired by descent. The burden is on the claiming spouse. **Tracing methods** — In re Marriage of Lucas (1980) 27 Cal.3d 808 held that when separate and community funds are commingled, the spouse must trace by direct evidence or by the "family expense" method (showing community expenses exceeded community income, so the acquisition came from separate funds). Obtain bank statements showing the deposit, the timing of the down payment, and that sufficient separate funds remained. **What if the account had substantial community deposits** — Tracing fails if she cannot isolate the source. If she used a mix of separate and community, she may get a Fam. Code § 2640 reimbursement for the separate contribution—dollar-for-dollar, no appreciation—but not a separate-property share of the house. Key authorities: Cal. Fam. Code §§ 760, 770, 2640 · In re Marriage of Lucas (1980) 27 Cal.3d 808

---

User: Client was pushed during an argument. She wants immediate protection. Ex parte DVRO or emergency protective order—what's the difference and what do I need to show?

Joogni: An ex parte restraining order under the Domestic Violence Prevention Act (DVPA) and an emergency protective order (EPO) under Penal Code § 136.2 serve different purposes. For an immediate crisis, an EPO—requested by law enforcement—lasts up to seven days and can be obtained 24/7. A DVRO ex parte under Fam. Code §§ 6320–6325 requires a court filing and a showing of "reasonable proof of a past act or threat of abuse." **Abuse under the DVPA** — Fam. Code § 6203 defines abuse to include assault, battery, and placing another in reasonable fear of bodily injury. A single push can suffice if it caused harm or fear. Document the incident, photos, witnesses, and any prior incidents. **What if the judge is skeptical** — Some departments require more than a single incident. If the ex parte is denied, request a short-cause hearing and consider an EPO in the meantime. An EPO does not require a court appearance by the petitioner—the officer requests it from an on-call judge. Have your client call 911 or go to the station if she needs protection before the courthouse opens. Key authorities: Cal. Fam. Code §§ 6200, 6203, 6320, 6325 · Cal. Penal Code § 136.2""",
    "ny": None,
    "tx": None,
    "fl": None,
    "il": None,
    "az": None,
    "nv": None,
    "wa": None,
    "or": None,
    "co": None,
}

FREE_CHAT_SYSTEM_PROMPT = """You are a helpful AI assistant. Answer questions clearly and concisely. Be conversational and helpful.

FORMATTING: Structure all responses for readability:
- Use markdown headings (# for main, ## for sections, ### for subsections)
- Use bullet lists (- or *) for unordered items
- Use numbered lists (1. 2. 3.) for sequential steps
- Add blank lines between sections; avoid dense text walls
- Keep paragraphs short (2-4 sentences max)"""

# Config per jurisdiction: implemented, domains for legal search, M365 tools enabled
JURISDICTION_CONFIG = {
    "ca": {
        "implemented": True,
        "display_name": "California",
        "domains": ["family_code", "evidence_code"],
        "has_tools": True,
    },
    "ny": {"implemented": False, "display_name": "New York", "domains": [], "has_tools": False},
    "tx": {"implemented": False, "display_name": "Texas", "domains": [], "has_tools": False},
    "fl": {"implemented": False, "display_name": "Florida", "domains": [], "has_tools": False},
    "il": {"implemented": False, "display_name": "Illinois", "domains": [], "has_tools": False},
    "az": {"implemented": False, "display_name": "Arizona", "domains": [], "has_tools": False},
    "nv": {"implemented": False, "display_name": "Nevada", "domains": [], "has_tools": False},
    "wa": {"implemented": False, "display_name": "Washington", "domains": [], "has_tools": False},
    "or": {"implemented": False, "display_name": "Oregon", "domains": [], "has_tools": False},
    "co": {"implemented": False, "display_name": "Colorado", "domains": [], "has_tools": False},
}


def resolve_jurisdiction(payload_value: str) -> tuple:
    """
    Returns (jurisdiction_id, config).
    If not implemented: config["implemented"] == False.
    """
    jid = JURISDICTION_NAME_TO_ID.get(payload_value) or JURISDICTION_NAME_TO_ID.get("California", "ca")
    config = JURISDICTION_CONFIG.get(jid, {"implemented": False, "domains": [], "has_tools": False})
    return jid, config


# Multi-domain legal search config (single index with domain field)
LEGAL_DOMAINS = [
    d.strip()
    for d in os.getenv("LEGAL_DOMAINS", "family_code,evidence_code").split(",")
    if d.strip()
]
LEGAL_MAX_DOMAINS = int(os.getenv("LEGAL_MAX_DOMAINS", "3"))
LEGAL_TOP_PER_DOMAIN = int(os.getenv("LEGAL_TOP_PER_DOMAIN", "3"))

# Upload / attachment settings
# Defaults raised to support uploads with multiple large PDFs; adjust via env vars.
MAX_UPLOAD_SIZE_MB = int(os.getenv("MAX_UPLOAD_SIZE_MB", "200"))
MAX_UPLOAD_BYTES = MAX_UPLOAD_SIZE_MB * 1024 * 1024
MAX_ATTACHMENT_FILES = int(os.getenv("MAX_ATTACHMENT_FILES", "50"))
# Optional: limit per-attachment size for chat flow; default 50MB to support 30-page high-res scans
_ATTACHMENT_MAX_MB = os.getenv("ATTACHMENT_MAX_SIZE_MB", "50")
ATTACHMENT_MAX_BYTES = int(_ATTACHMENT_MAX_MB) * 1024 * 1024
ATTACHMENT_MAX_CHUNKS = int(os.getenv("ATTACHMENT_MAX_CHUNKS", "50"))
EMBEDDING_DEPLOYMENT = os.getenv("AZURE_OPENAI_EMBEDDING_DEPLOYMENT", "text-embedding-3-large")
VECTOR_DIMENSION = int(os.getenv("VECTOR_DIMENSION", "3072"))  # text-embedding-3-large; must match Azure Search index contentVector
PDF_OCR_MAX_PAGES = int(os.getenv("PDF_OCR_MAX_PAGES", "30"))
PDF_OCR_DPI = int(os.getenv("PDF_OCR_DPI", "150"))
PDF_OCR_LANG = os.getenv("PDF_OCR_LANG", "eng")
# Below this token count, send full extracted text to LLM (skip RAG indexing)
ATTACHMENT_FULL_TEXT_MAX_TOKENS = int(os.getenv("ATTACHMENT_FULL_TEXT_MAX_TOKENS", "4000"))

deployment_name = os.getenv("AZURE_OPENAI_EMBEDDING_DEPLOYMENT")
print("AZURE_OPENAI_EMBEDDING_DEPLOYMENT present:", bool(deployment_name))

def _search_one_domain_sync(
    search_client: SearchClient,
    query: str,
    query_vector: Optional[List[float]],
    dom: str,
    top_per_domain: int,
):
    """Sync helper: search one domain and return (dom, chunks). Used for parallel execution."""
    from azure.search.documents.models import VectorizedQuery

    kwargs = {"search_text": query, "top": top_per_domain}
    kwargs["filter"] = f"domain eq '{dom}'"
    if query_vector:
        kwargs["vector_queries"] = [
            VectorizedQuery(
                vector=query_vector,
                k_nearest_neighbors=top_per_domain,
                fields="contentVector",
            )
        ]
    results = search_client.search(**kwargs)
    chunks = []
    for r in results:
        content = (r.get("content") or "").strip()
        meta = (r.get("metadata") or "").strip()
        source = r.get("filepath") or r.get("source") or r.get("url") or ""
        if content:
            chunks.append(f"[DOMAIN: {dom}] [Source: {source}]\n{content}\nMeta: {meta}")
    return (dom, chunks)


def _search_fallback_sync(
    search_client: SearchClient,
    query: str,
    query_vector: Optional[List[float]],
    top_per_domain: int,
):
    """Sync helper: fallback search when no domain has hits."""
    from azure.search.documents.models import VectorizedQuery

    kwargs = {"search_text": query, "top": top_per_domain}
    if query_vector:
        kwargs["vector_queries"] = [
            VectorizedQuery(
                vector=query_vector,
                k_nearest_neighbors=top_per_domain,
                fields="contentVector",
            )
        ]
    results = search_client.search(**kwargs)
    chunks = []
    for r in results:
        content = (r.get("content") or "").strip()
        meta = (r.get("metadata") or "").strip()
        source = r.get("filepath") or r.get("source") or r.get("url") or ""
        if content:
            chunks.append(
                f"[DOMAIN: unknown] [Source: {source}]\n{content}\nMeta: {meta}"
            )
    return chunks


def run_multi_domain_search(
    search_client: SearchClient,
    query: str,
    domains: List[str],
    max_domains: int,
    top_per_domain: int,
    embed_client: Optional["AzureOpenAI"] = None,
):
    """
    Hybrid (vector + keyword) or keyword-only multi-domain search.
    Sync version for backwards compatibility.
    """
    from azure.search.documents.models import VectorizedQuery

    context_chunks = []
    domains_with_hits = []
    candidates = [d.strip() for d in domains if d.strip()][:max_domains]

    query_vector = None
    if embed_client and EMBEDDING_DEPLOYMENT and query.strip():
        try:
            emb = embed_client.embeddings.create(
                model=EMBEDDING_DEPLOYMENT,
                input=[query.strip()[:8000]],
                dimensions=VECTOR_DIMENSION,
            )
            query_vector = emb.data[0].embedding
        except Exception as e:
            print(f"Embedding error (fallback to keyword only): {e}", file=sys.stderr)

    for dom in candidates:
        try:
            _, chunks = _search_one_domain_sync(
                search_client, query, query_vector, dom, top_per_domain
            )
            context_chunks.extend(chunks)
            if chunks:
                domains_with_hits.append(dom)
        except Exception as e:
            print(f"Search error for domain {dom}: {e}", file=sys.stderr)

    if not context_chunks:
        try:
            fallback_chunks = _search_fallback_sync(
                search_client, query, query_vector, top_per_domain
            )
            context_chunks.extend(fallback_chunks)
            if fallback_chunks:
                domains_with_hits.append("unknown")
        except Exception as e:
            print(f"Fallback search error: {e}", file=sys.stderr)

    return context_chunks, domains_with_hits


async def run_multi_domain_search_async(
    search_client: SearchClient,
    query: str,
    domains: List[str],
    max_domains: int,
    top_per_domain: int,
    embed_client: Optional["AzureOpenAI"] = None,
):
    """
    Async version: embedding and domain searches run via thread pool (no event loop blocking).
    Domain searches run in parallel; results merged in deterministic domain order for quality.
    """
    candidates = [d.strip() for d in domains if d.strip()][:max_domains]
    query_vector = None

    if embed_client and EMBEDDING_DEPLOYMENT and query.strip():
        try:
            emb = await asyncio.to_thread(
                embed_client.embeddings.create,
                model=EMBEDDING_DEPLOYMENT,
                input=[query.strip()[:8000]],
                dimensions=VECTOR_DIMENSION,
            )
            query_vector = emb.data[0].embedding
        except Exception as e:
            print(f"Embedding error (fallback to keyword only): {e}", file=sys.stderr)

    # Run domain searches in parallel; merge in candidates order (deterministic)
    tasks = [
        asyncio.to_thread(
            _search_one_domain_sync,
            search_client,
            query,
            query_vector,
            dom,
            top_per_domain,
        )
        for dom in candidates
    ]
    domain_results = await asyncio.gather(*tasks, return_exceptions=True)

    context_chunks = []
    domains_with_hits = []
    for i, res in enumerate(domain_results):
        if isinstance(res, Exception):
            print(f"Search error for domain {candidates[i]}: {res}", file=sys.stderr)
            continue
        dom, chunks = res
        context_chunks.extend(chunks)
        if chunks:
            domains_with_hits.append(dom)

    if not context_chunks:
        try:
            fallback_chunks = await asyncio.to_thread(
                _search_fallback_sync,
                search_client,
                query,
                query_vector,
                top_per_domain,
            )
            context_chunks.extend(fallback_chunks)
            if fallback_chunks:
                domains_with_hits.append("unknown")
        except Exception as e:
            print(f"Fallback search error: {e}", file=sys.stderr)

    return context_chunks, domains_with_hits


def _run_user_attachment_search_sync(
    search_client: SearchClient,
    query: str,
    blob_names: List[str],
    top_k: int,
    embed_client: Optional["AzureOpenAI"],
) -> List[str]:
    """Sync implementation of user attachment search."""
    if not blob_names:
        return []
    from azure.search.documents.models import VectorizedQuery

    query_vector = None
    if embed_client and EMBEDDING_DEPLOYMENT and query.strip():
        try:
            emb = embed_client.embeddings.create(
                model=EMBEDDING_DEPLOYMENT,
                input=[query.strip()[:8000]],
                dimensions=VECTOR_DIMENSION,
            )
            query_vector = emb.data[0].embedding
        except Exception as e:
            print(f"[UserAttachment] Embedding error: {e}", file=sys.stderr)
    safe_sources = [s.replace("'", "''") for s in blob_names if s]
    if not safe_sources:
        return []
    source_filter = " or ".join(f"source eq '{s}'" for s in safe_sources)
    filter_str = f"domain eq 'user_attachment' and ({source_filter})"
    kwargs = {"search_text": query, "top": top_k, "filter": filter_str}
    if query_vector:
        kwargs["vector_queries"] = [
            VectorizedQuery(
                vector=query_vector, k_nearest_neighbors=top_k, fields="contentVector"
            )
        ]
    context_chunks = []
    try:
        results = search_client.search(**kwargs)
        for r in results:
            content = (r.get("content") or "").strip()
            source = r.get("filepath") or r.get("source") or ""
            title = r.get("title") or source
            if content:
                context_chunks.append(f"[Attachment: {title}]\n{content}")
    except Exception as e:
        print(f"[UserAttachment] Search error: {e}", file=sys.stderr)
    return context_chunks


def run_user_attachment_search(
    search_client: SearchClient,
    query: str,
    blob_names: List[str],
    top_k: int = 6,
    embed_client: Optional["AzureOpenAI"] = None,
) -> List[str]:
    """Search user_attachment domain for the given blob_names. Sync version."""
    return _run_user_attachment_search_sync(
        search_client, query, blob_names, top_k, embed_client
    )


async def run_user_attachment_search_async(
    search_client: SearchClient,
    query: str,
    blob_names: List[str],
    top_k: int = 6,
    embed_client: Optional["AzureOpenAI"] = None,
) -> List[str]:
    """Async version: runs in thread pool to avoid blocking event loop."""
    return await asyncio.to_thread(
        _run_user_attachment_search_sync,
        search_client,
        query,
        blob_names,
        top_k,
        embed_client,
    )


# Define tools for function calling
TOOLS = [
    {
        "type": "function",
        "function": {
            "name": "search_emails",
            "description": "Search the user's Outlook emails. Use this when the user asks about emails, messages, correspondence, or communications from specific people or about specific topics.",
            "parameters": {
                "type": "object",
                "properties": {
                    "query": {
                        "type": "string",
                        "description": "Search query - can be a person's name, email address, subject, or keywords"
                    },
                    "limit": {
                        "type": "integer",
                        "description": "Maximum number of emails to return (default 10)",
                        "default": 10
                    }
                },
                "required": ["query"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "get_email_content",
            "description": "Get the full content of a specific email by its ID. Use this after search_emails to read the full body of an email.",
            "parameters": {
                "type": "object",
                "properties": {
                    "message_id": {
                        "type": "string",
                        "description": "The ID of the email message to retrieve"
                    }
                },
                "required": ["message_id"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "search_calendar",
            "description": "Search the user's calendar for meetings, appointments, hearings, or events. Use this when the user asks about their schedule, meetings, hearings, or calendar.",
            "parameters": {
                "type": "object",
                "properties": {
                    "query": {
                        "type": "string",
                        "description": "Optional search term to filter events by subject or location. Leave empty to get all upcoming events.",
                        "default": ""
                    },
                    "days_ahead": {
                        "type": "integer",
                        "description": "Number of days ahead to search (default 7 for a week)",
                        "default": 7
                    }
                },
                "required": []
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "search_files",
            "description": "Search the user's OneDrive files. Use this when the user asks about documents, files, or wants to find specific files.",
            "parameters": {
                "type": "object",
                "properties": {
                    "query": {
                        "type": "string",
                        "description": "Search query - filename, content keywords, or file type"
                    }
                },
                "required": ["query"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "get_todays_events",
            "description": "Get all of today's calendar events. Use this when the user asks 'do I have meetings today', 'what's on my calendar today', or similar.",
            "parameters": {
                "type": "object",
                "properties": {},
                "required": []
            }
        }
    }
]


def get_tools(jurisdiction_id: str, config: dict, graph_token_available: bool) -> Optional[List]:
    """Returns tools for jurisdiction. Only CA has functional tools at the moment."""
    if not config.get("implemented") or not config.get("has_tools"):
        return None
    return TOOLS if graph_token_available else None


def get_user_info(request: Request) -> dict:
    """Extract user info from Azure Easy Auth headers."""
    principal = request.headers.get("X-MS-CLIENT-PRINCIPAL")
    if principal:
        try:
            decoded = base64.b64decode(principal)
            return json.loads(decoded)
        except:
            pass
    return None


def is_authenticated(request: Request) -> bool:
    """Check if user is authenticated via Azure Easy Auth."""
    principal = request.headers.get("X-MS-CLIENT-PRINCIPAL")
    id_token = request.headers.get("X-MS-TOKEN-AAD-ID-TOKEN")
    return bool(principal or id_token)


def get_graph_token(request: Request) -> Optional[str]:
    """Get Microsoft Graph access token from Easy Auth."""
    token = request.headers.get("X-MS-TOKEN-AAD-ACCESS-TOKEN")
    principal = request.headers.get("X-MS-CLIENT-PRINCIPAL")
    id_token = request.headers.get("X-MS-TOKEN-AAD-ID-TOKEN")
    expires_on = request.headers.get("X-MS-TOKEN-AAD-EXPIRES-ON")
    print(
        "Graph token presence: "
        f"access={bool(token)} len={len(token) if token else 0} "
        f"principal={bool(principal)} id_token={bool(id_token)} "
        f"expires_on={expires_on}",
        file=sys.stderr,
    )
    if token:
        print(f"Graph token prefix: {token[:10]}...", file=sys.stderr)
    return token or None


def _get_token_expiry(request: Request) -> Optional[int]:
    """Return token expiry (epoch seconds) from Easy Auth header."""
    expires_on = request.headers.get("X-MS-TOKEN-AAD-EXPIRES-ON")
    if not expires_on:
        return None
    try:
        return int(expires_on)
    except Exception:
        # Some environments send ISO strings like 2026-01-26T23:07:48.0666466Z
        try:
            iso = expires_on.replace("Z", "+00:00")
            return int(datetime.fromisoformat(iso).timestamp())
        except Exception:
            return None


def _token_near_expiry(expiry: Optional[int], skew_seconds: int = 300) -> bool:
    """True if expiry is within skew window."""
    if not expiry:
        return False
    now = int(time.time())
    return expiry - now <= skew_seconds


async def refresh_graph_token(request: Request) -> Optional[str]:
    """
    Attempt to refresh Easy Auth tokens via /.auth/refresh using the caller's cookies.
    Returns a new access token if present in the response.
    """
    cookies = request.headers.get("cookie")
    if not cookies:
        print("[Graph token] No cookies on request; cannot refresh.", file=sys.stderr)
        return None

    refresh_url = str(request.base_url).rstrip("/") + "/.auth/refresh"
    try:
        async with httpx.AsyncClient(follow_redirects=False) as client:
            resp = await client.get(
                refresh_url,
                headers={"cookie": cookies},
                timeout=10.0,
            )
    except Exception as err:
        print(f"[Graph token] Refresh call failed: {err}", file=sys.stderr)
        return None

    if resp.status_code not in (200, 302):
        print(
            f"[Graph token] Refresh failed status={resp.status_code} body={resp.text[:200]}",
            file=sys.stderr,
        )
        return None

    try:
        data = resp.json()
    except Exception as err:
        print(f"[Graph token] Refresh JSON parse error: {err}", file=sys.stderr)
        return None

    if not isinstance(data, dict):
        print("[Graph token] Refresh response not a dict; cannot extract token.", file=sys.stderr)
        return None

    new_token = (
        data.get("access_token")
        or data.get("access_token_v2")
        or data.get("authenticationToken")
    )
    print(
        f"[Graph token] Refresh success token_present={bool(new_token)} "
        f"expires_on={data.get('expires_on')}",
        file=sys.stderr,
    )
    return new_token


async def get_valid_graph_token(request: Request, refresh_skew_seconds: int = 300) -> Optional[str]:
    """
    Return a usable Graph token, attempting refresh if the Easy Auth token is close to expiry.
    """
    token = get_graph_token(request)
    expiry = _get_token_expiry(request)
    if token and not _token_near_expiry(expiry, refresh_skew_seconds):
        return token

    refreshed = await refresh_graph_token(request)
    if refreshed:
        return refreshed

    # Fall back to whatever we had (may be None / expired)
    return token


def _get_blob_service_client() -> BlobServiceClient:
    """
    Return a BlobServiceClient using connection string if present, otherwise Managed Identity.
    """
    connection_string = os.getenv("AZURE_STORAGE_CONNECTION_STRING")
    account_name = os.getenv("AZURE_STORAGE_ACCOUNT_NAME", "glgaistorage")

    if connection_string:
        return BlobServiceClient.from_connection_string(connection_string)

    credential = DefaultAzureCredential()
    account_url = f"https://{account_name}.blob.core.windows.net"
    return BlobServiceClient(account_url, credential=credential)


def _ensure_container(blob_service: BlobServiceClient, container: str):
    container_client = blob_service.get_container_client(container)
    try:
        container_client.create_container()
    except Exception:
        # likely already exists
        pass


def _generate_sas_url(blob_service: BlobServiceClient, container: str, blob_name: str, expiry_hours: int = 1) -> Tuple[str, str]:
    """
    Generate a SAS URL for uploading a blob. Supports both account key and user delegation key.
    Returns (upload_url, blob_name).
    """
    expiry = datetime.utcnow() + timedelta(hours=expiry_hours)
    account_name = blob_service.account_name

    sas_token = None
    # Try account key first (connection string case)
    credential = getattr(blob_service, "credential", None)
    account_key = getattr(credential, "account_key", None)

    if account_key:
        sas_token = generate_blob_sas(
            account_name=account_name,
            container_name=container,
            blob_name=blob_name,
            account_key=account_key,
            permission=BlobSasPermissions(write=True, create=True),
            expiry=expiry,
        )
    else:
        # Managed Identity path: use user delegation key
        delegation_key = blob_service.get_user_delegation_key(datetime.utcnow(), expiry)
        sas_token = generate_blob_sas(
            account_name=account_name,
            container_name=container,
            blob_name=blob_name,
            user_delegation_key=delegation_key,
            permission=BlobSasPermissions(write=True, create=True),
            expiry=expiry,
        )

    blob_client = blob_service.get_blob_client(container=container, blob=blob_name)
    upload_url = f"{blob_client.url}?{sas_token}"
    return upload_url, blob_name


def _get_user_blob_prefix(request: Request) -> str:
    """Return a stable prefix for the current user's blobs (isolation by user)."""
    user_info = get_user_info(request)
    uid = ""
    if user_info:
        uid = (
            user_info.get("userId")
            or user_info.get("user_id")
            or user_info.get("email")
            or ""
        )
        if not uid and isinstance(user_info.get("userDetails"), str):
            uid = user_info["userDetails"]
    raw = (uid or "anonymous").strip().lower()
    h = hashlib.sha256(raw.encode("utf-8")).hexdigest()[:16]
    return f"u_{h}"


def _safe_blob_name(filename: str, user_prefix: str = "") -> str:
    base = os.path.basename(filename)
    base = re.sub(r"[^A-Za-z0-9._-]", "_", base)
    name = f"{uuid.uuid4()}-{base}"
    if user_prefix:
        return f"{user_prefix}/{name}"
    return name


def _iter_pdf_text_pages(file_bytes: bytes):
    """Generator yielding text from each PDF page (avoids building full list in memory)."""
    reader = PdfReader(io.BytesIO(file_bytes))
    for page in reader.pages:
        yield page.extract_text() or ""


def _iter_pdf_ocr_pages(file_bytes: bytes, max_pages: int):
    """Generator yielding OCR text from each PDF page (avoids accumulating ocr_chunks list)."""
    from pdf2image import convert_from_bytes  # type: ignore
    import pytesseract  # type: ignore

    for page_num in range(1, max_pages + 1):
        try:
            images = convert_from_bytes(
                file_bytes,
                dpi=PDF_OCR_DPI,
                first_page=page_num,
                last_page=page_num,
            )
        except Exception as conv_err:
            print(f"OCR conversion failed on page {page_num}: {conv_err}", file=sys.stderr)
            continue

        for img in images:
            try:
                txt = pytesseract.image_to_string(img, lang=PDF_OCR_LANG) or ""
            except Exception as ocr_err:
                print(f"OCR failed on page {page_num}: {ocr_err}", file=sys.stderr)
                continue
            finally:
                try:
                    img.close()
                except Exception:
                    pass
            if txt.strip():
                yield txt.strip()

        del images
        gc.collect()


_ATTACHMENT_ALLOWED_EXT = (".pdf", ".docx")


def _is_supported_attachment(filename: str) -> bool:
    """Return True if filename has an allowed extension for attachments."""
    fn = (filename or "").lower().strip()
    return any(fn.endswith(ext) for ext in _ATTACHMENT_ALLOWED_EXT)


def _extract_text_from_docx(file_bytes: bytes, filename: str) -> str:
    """Extract text from Word .docx file."""
    try:
        from docx import Document

        doc = Document(io.BytesIO(file_bytes))
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text.strip())
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text.strip())
        text = "\n\n".join(parts).strip()
        if not text:
            raise HTTPException(
                status_code=400,
                detail=f"Word document {filename} appears to be empty or has no extractable text.",
            )
        return text
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to read Word document {filename}: {e}")


def _extract_text_from_pdf(file_bytes: bytes, filename: str) -> str:
    """Extract text from PDF using page-by-page generators to reduce memory usage."""
    try:
        content = "\n".join(_iter_pdf_text_pages(file_bytes)).strip()
        if content:
            return content

        reader = PdfReader(io.BytesIO(file_bytes))
        page_count = len(reader.pages)
        if not page_count:
            raise HTTPException(status_code=400, detail=f"PDF {filename} is empty or has no readable pages.")

        max_pages = min(page_count, PDF_OCR_MAX_PAGES)
        print(f"[PDF] OCR needed for {filename}: {page_count} pages, processing {max_pages}", file=sys.stderr)
        if page_count > PDF_OCR_MAX_PAGES:
            raise HTTPException(
                status_code=400,
                detail=(
                    f"PDF {filename} has {page_count} pages; OCR limit is {PDF_OCR_MAX_PAGES}. "
                    "Upload a smaller file or one with embedded text."
                ),
            )

        try:
            from pdf2image import convert_from_bytes  # type: ignore
            import pytesseract  # type: ignore
        except Exception as import_err:
            raise HTTPException(
                status_code=500,
                detail=(
                    f"OCR unavailable for {filename}: install poppler + tesseract on the system and "
                    f"pdf2image/pytesseract libraries. Error: {import_err}"
                ),
            ) from import_err

        ocr_content = "\n".join(_iter_pdf_ocr_pages(file_bytes, max_pages)).strip()
        if ocr_content:
            return ocr_content

        raise HTTPException(
            status_code=400,
            detail=(
                f"PDF {filename} appears to be scanned or redacted and OCR found no text. "
                "Upload a clearer copy or one with embedded OCR text."
            ),
        )

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to read PDF {filename}: {e}")


def _extract_text_from_file(file_bytes: bytes, filename: str) -> str:
    """Extract text from supported file (PDF or Word .docx)."""
    fn = (filename or "").lower()
    if fn.endswith(".pdf"):
        return _extract_text_from_pdf(file_bytes, filename)
    if fn.endswith(".docx"):
        return _extract_text_from_docx(file_bytes, filename)
    raise HTTPException(
        status_code=400,
        detail=f"Unsupported file type: {filename}. Only PDF and Word (.docx) are allowed.",
    )


_tiktoken_encoder: Optional[object] = "unset"  # "unset" | encoder | None (failed, use fallback)


def _get_tiktoken_encoder():
    """Lazy-load tiktoken encoder. Returns None if network unavailable (offline)."""
    global _tiktoken_encoder
    if _tiktoken_encoder != "unset":
        return _tiktoken_encoder if _tiktoken_encoder is not None else None
    try:
        _tiktoken_encoder = tiktoken.get_encoding("cl100k_base")
        return _tiktoken_encoder
    except Exception as e:
        print(f"[Chunk] tiktoken unavailable ({e}), using char-based fallback", file=sys.stderr)
        _tiktoken_encoder = None
        return None


def _chunk_text(text: str, max_tokens: int = 350, overlap: int = 60, max_chunks: int = 200) -> List[str]:
    """
    Tries tiktoken first (exact token counting). Falls back to char-based chunking
    when tiktoken cannot load (e.g. offline / network unreachable).
    """
    encoder = _get_tiktoken_encoder()
    if encoder is not None:
        tokens = encoder.encode(text)
        chunks = []
        start = 0
        while start < len(tokens) and len(chunks) < max_chunks:
            end = min(start + max_tokens, len(tokens))
            chunk_tokens = tokens[start:end]
            chunks.append(encoder.decode(chunk_tokens))
            if end == len(tokens):
                break
            start = end - overlap
        return chunks

    # Fallback: ~4 chars/token, 350 tokens ~1400 chars, 60 overlap ~240 chars
    chunk_size = max_tokens * 4
    overlap_chars = overlap * 4
    chunks = []
    start = 0
    while start < len(text) and len(chunks) < max_chunks:
        end = min(start + chunk_size, len(text))
        chunks.append(text[start:end])
        if end == len(text):
            break
        start = end - overlap_chars
    return chunks


def _count_tokens(text: str) -> int:
    """Return token count using tiktoken. Falls back to char/4 estimate if unavailable."""
    encoder = _get_tiktoken_encoder()
    if encoder is not None:
        return len(encoder.encode(text))
    return len(text) // 4


def _cosine_similarity(a: List[float], b: List[float]) -> float:
    denom = math.sqrt(sum(x * x for x in a)) * math.sqrt(sum(y * y for y in b))
    if denom == 0:
        return 0.0
    return sum(x * y for x, y in zip(a, b)) / denom


EMBED_BATCH_SIZE = 24


async def _fetch_full_text_context(blobs: List[dict]) -> List[str]:
    """
    Download and extract full text from blobs marked use_full_text (small docs that skipped RAG).
    Returns list of context strings: [f"[Attachment: {filename}]\n{text}", ...]
    """
    if not blobs:
        return []
    storage = LegalDocsStorage()
    upload_container = os.getenv("UPLOAD_CONTAINER") or os.getenv("AZURE_STORAGE_CONTAINER") or "legal-docs-raw"
    sections: List[str] = []
    for blob in blobs:
        blob_name = blob.get("blob_name")
        original_filename = blob.get("original_filename", "file.pdf")
        if not blob_name:
            continue
        try:
            file_bytes = await asyncio.to_thread(storage.download_file, blob_name, container=upload_container)
            if not file_bytes or len(file_bytes) > ATTACHMENT_MAX_BYTES:
                continue
            if not _is_supported_attachment(original_filename):
                continue
            text = await asyncio.to_thread(_extract_text_from_file, file_bytes, original_filename)
            del file_bytes
            gc.collect()
            if not text or not text.strip():
                continue
            sections.append(f"[Attachment: {original_filename}]\n{text.strip()}")
        except Exception as e:
            print(f"[Attachment] Full-text fetch failed for {original_filename}: {e}", file=sys.stderr)
    return sections


async def _build_attachment_context(blobs: List[dict], query: str, client: AzureOpenAI) -> List[str]:
    """
    Download, chunk, embed attachments and return top chunks as context strings.
    Uses batching (24 chunks at a time), thread pool for blocking I/O, and explicit memory release to avoid OOM.
    """
    if not blobs:
        return []
    if not EMBEDDING_DEPLOYMENT:
        raise HTTPException(status_code=500, detail="Embedding deployment not configured")

    if len(blobs) > MAX_ATTACHMENT_FILES:
        raise HTTPException(status_code=400, detail=f"Too many files. Max {MAX_ATTACHMENT_FILES}.")

    storage = LegalDocsStorage()
    upload_container = os.getenv("UPLOAD_CONTAINER") or os.getenv("AZURE_STORAGE_CONTAINER") or "legal-docs-raw"
    attachment_sections: List[str] = []
    top_chunks_per_file = 3

    # Embed query once (run in thread to avoid blocking event loop)
    safe_query = str(query or "legal question")
    t0 = time.perf_counter()
    query_embed_resp = await asyncio.to_thread(
        lambda: client.embeddings.create(
            model=EMBEDDING_DEPLOYMENT,
            input=[safe_query],
            dimensions=VECTOR_DIMENSION,
        )
    )
    query_vec = query_embed_resp.data[0].embedding
    print(f"[Attachment] Query embed done in {time.perf_counter() - t0:.2f}s", file=sys.stderr)

    for blob in blobs:
        blob_name = blob.get("blob_name")
        original_filename = blob.get("original_filename", "file.pdf")
        if not blob_name:
            continue
        file_bytes = b""
        text = ""
        chunk_texts: List[str] = []
        try:
            t_dl = time.perf_counter()
            file_bytes = await asyncio.to_thread(storage.download_file, blob_name, container=upload_container)
            print(f"[Attachment] Downloaded {original_filename}: {len(file_bytes)} bytes in {time.perf_counter() - t_dl:.2f}s", file=sys.stderr)

            if len(file_bytes) > ATTACHMENT_MAX_BYTES:
                raise HTTPException(status_code=400, detail=f"File {original_filename} exceeds {_ATTACHMENT_MAX_MB}MB limit. Use a smaller file or increase ATTACHMENT_MAX_SIZE_MB.")

            if _is_supported_attachment(original_filename):
                t_ext = time.perf_counter()
                text = await asyncio.to_thread(_extract_text_from_file, file_bytes, original_filename)
                print(f"[Attachment] Extract done for {original_filename} in {time.perf_counter() - t_ext:.2f}s", file=sys.stderr)
            else:
                raise HTTPException(
                    status_code=400,
                    detail=f"Unsupported file type: {original_filename}. Only PDF and Word (.docx) are allowed.",
                )

            del file_bytes
            gc.collect()

            chunk_texts = _chunk_text(text, max_chunks=ATTACHMENT_MAX_CHUNKS)
            del text
            gc.collect()
            print(f"[Attachment] Chunked {original_filename}: {len(chunk_texts)} chunks", file=sys.stderr)
        except Exception:
            raise

        if not chunk_texts:
            continue

        # Embed in batches; keep top 3 per batch, then merge for global top 3 (run in thread)
        batch_top_scores: List[Tuple[float, str]] = []
        n_batches = (len(chunk_texts) + EMBED_BATCH_SIZE - 1) // EMBED_BATCH_SIZE
        for i in range(0, len(chunk_texts), EMBED_BATCH_SIZE):
            batch = [str(c) for c in chunk_texts[i : i + EMBED_BATCH_SIZE]]
            batch_idx = i // EMBED_BATCH_SIZE + 1
            t_emb = time.perf_counter()
            embed_resp = await asyncio.to_thread(
                lambda b=batch: client.embeddings.create(
                    model=EMBEDDING_DEPLOYMENT,
                    input=b,
                    dimensions=VECTOR_DIMENSION,
                )
            )
            for chunk_text, emb in zip(batch, embed_resp.data):
                score = _cosine_similarity(query_vec, emb.embedding)
                batch_top_scores.append((score, chunk_text))
            del embed_resp
            gc.collect()
            print(f"[Attachment] Embed batch {batch_idx}/{n_batches} ({len(batch)} chunks) in {time.perf_counter() - t_emb:.2f}s", file=sys.stderr)

        batch_top_scores.sort(key=lambda x: x[0], reverse=True)
        top_chunks = batch_top_scores[:top_chunks_per_file]
        del batch_top_scores
        del chunk_texts
        gc.collect()

        if not top_chunks:
            continue

        section_lines = [f"[Attachment: {original_filename}]"]
        for idx, (_, chunk_text) in enumerate(top_chunks, start=1):
            section_lines.append(f"{idx}. {chunk_text}")
        attachment_sections.append("\n".join(section_lines))

    return attachment_sections


def _check_search_connectivity(search_service: str, label: str = "Attachment") -> Optional[str]:
    """Test connectivity to Azure Search. Returns None if OK, error message if failed."""
    import socket
    import threading
    host = f"{search_service}.search.windows.net"
    port = 443
    thread_name = threading.current_thread().name
    try:
        print(f"[{label}] Connectivity check (thread={thread_name}): resolving {host}...", file=sys.stderr)
        ip = socket.gethostbyname(host)
        print(f"[{label}] Resolved {host} -> {ip}", file=sys.stderr)
        sock = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
        sock.settimeout(10)
        sock.connect((host, port))
        sock.close()
        print(f"[{label}] TCP connection to {host}:{port} OK (thread={thread_name})", file=sys.stderr)
        return None
    except socket.gaierror as e:
        err = f"DNS/resolution failed for {host}: {e}"
        print(f"[{label}] {err}", file=sys.stderr)
        return err
    except OSError as e:
        err = f"Connection to {host}:{port} failed (errno={getattr(e, 'errno', '?')}): {e}"
        print(f"[{label}] {err}", file=sys.stderr)
        return err


def process_attachment_job(job_id: str, partition_key: str, blob_name: str, original_filename: str):
    """
    Background task: download blob, extract PDF, chunk, embed, index to Azure Search.
    Updates job status to completed or failed.
    """
    try:
        update_job_status(partition_key, job_id, "processing")
        search_service = os.getenv("AZURE_SEARCH_SERVICE")
        if search_service:
            conn_err = _check_search_connectivity(search_service, label="BackgroundTask")
            if conn_err:
                print(f"[Attachment] Pre-flight failed: {conn_err}", file=sys.stderr)
                update_job_status(partition_key, job_id, "failed", conn_err)
                return
        storage = LegalDocsStorage()
        upload_container = os.getenv("UPLOAD_CONTAINER") or os.getenv("AZURE_STORAGE_CONTAINER") or "legal-docs-raw"

        file_bytes = storage.download_file(blob_name, container=upload_container)
        if not file_bytes or len(file_bytes) > ATTACHMENT_MAX_BYTES:
            update_job_status(partition_key, job_id, "failed", "File too large or empty")
            return
        if not _is_supported_attachment(original_filename):
            update_job_status(partition_key, job_id, "failed", f"Unsupported file type. Only PDF and Word (.docx) are allowed.")
            return
        text = _extract_text_from_file(file_bytes, original_filename)
        del file_bytes
        gc.collect()
        token_count = _count_tokens(text)
        if token_count <= ATTACHMENT_FULL_TEXT_MAX_TOKENS:
            update_job_status(partition_key, job_id, "completed", use_full_text=True)
            print(f"[Attachment] Job {job_id} completed (full-text): {blob_name} -> {token_count} tokens", file=sys.stderr)
            return
        chunk_texts = _chunk_text(text, max_chunks=ATTACHMENT_MAX_CHUNKS)
        del text
        gc.collect()
        if not chunk_texts:
            update_job_status(partition_key, job_id, "completed")  # No text, nothing to index
            return

        key = os.getenv("AZURE_OPENAI_KEY")
        endpoint = os.getenv("AZURE_OPENAI_ENDPOINT")
        search_service = os.getenv("AZURE_SEARCH_SERVICE")
        search_index = os.getenv("AZURE_SEARCH_INDEX")
        search_key = os.getenv("AZURE_SEARCH_KEY")
        if not all([key, endpoint, search_service, search_index, search_key]):
            update_job_status(partition_key, job_id, "failed", "Azure config missing")
            return

        client = AzureOpenAI(api_key=key, azure_endpoint=endpoint, api_version=os.getenv("AZURE_OPENAI_API_VERSION", "2023-05-15"))
        search_client = SearchClient(
            f"https://{search_service}.search.windows.net",
            search_index,
            AzureKeyCredential(search_key),
        )
        blob_stem = re.sub(r"[^a-zA-Z0-9_-]", "_", blob_name.replace("/", "_").replace(".pdf", ""))[:64]
        docs_to_upload = []
        for i in range(0, len(chunk_texts), EMBED_BATCH_SIZE):
            batch = [str(c) for c in chunk_texts[i : i + EMBED_BATCH_SIZE]]
            emb_resp = client.embeddings.create(
                model=EMBEDDING_DEPLOYMENT,
                input=batch,
                dimensions=VECTOR_DIMENSION,
            )
            for idx, (chunk_text, emb_obj) in enumerate(zip(batch, emb_resp.data)):
                chunk_id = f"{blob_stem}-{i + idx}"
                docs_to_upload.append({
                    "@search.action": "upload",
                    "id": chunk_id,
                    "content": chunk_text,
                    "domain": "user_attachment",
                    "source": blob_name,
                    "filepath": blob_name,
                    "title": original_filename,
                    "metadata": json.dumps({"original_filename": original_filename}),
                    "contentVector": emb_obj.embedding,
                })
            del emb_resp
            gc.collect()
        if docs_to_upload:
            batch_size = 50
            for j in range(0, len(docs_to_upload), batch_size):
                search_client.upload_documents(documents=docs_to_upload[j : j + batch_size])
        update_job_status(partition_key, job_id, "completed")
        print(f"[Attachment] Job {job_id} completed: {blob_name} -> {len(docs_to_upload)} chunks", file=sys.stderr)
    except Exception as e:
        import traceback
        traceback.print_exc()
        update_job_status(partition_key, job_id, "failed", str(e))
        print(f"[Attachment] Job {job_id} failed: {e}", file=sys.stderr)


# ============== Tool Execution Functions ==============

async def execute_search_emails(token: str, query: str, limit: int = 10, request: Optional[Request] = None) -> dict:
    """Execute email search via Microsoft Graph."""
    try:
        async with httpx.AsyncClient() as client:
            url = f"https://graph.microsoft.com/v1.0/me/messages?$search=\"{query}\"&$top={limit}&$select=id,subject,from,receivedDateTime,bodyPreview,hasAttachments"
            print(f"[Graph] search_emails url={url}", file=sys.stderr)
            
            headers = {"Authorization": f"Bearer {token}"}
            response = await client.get(url, headers=headers, timeout=30.0)
            
            # If token expired, try one refresh + retry (best effort)
            if response.status_code == 401 and request:
                print("[Graph] search_emails 401; attempting token refresh + retry", file=sys.stderr)
                refreshed = await refresh_graph_token(request)
                if refreshed:
                    headers = {"Authorization": f"Bearer {refreshed}"}
                    response = await client.get(url, headers=headers, timeout=30.0)
                    token = refreshed  # use refreshed token for parsing below
                    print(f"[Graph] search_emails retry status={response.status_code}", file=sys.stderr)
            
            if response.status_code != 200:
                print(f"[Graph] search_emails failed status={response.status_code} body={response.text[:300]}", file=sys.stderr)
                return {"error": f"Failed to search emails: {response.status_code}"}
            
            data = response.json()
            emails = []
            for msg in data.get("value", []):
                emails.append({
                    "id": msg.get("id"),
                    "subject": msg.get("subject", "(No subject)"),
                    "from": msg.get("from", {}).get("emailAddress", {}).get("name", "Unknown"),
                    "from_email": msg.get("from", {}).get("emailAddress", {}).get("address", ""),
                    "date": msg.get("receivedDateTime", ""),
                    "preview": msg.get("bodyPreview", "")[:200],
                    "hasAttachments": msg.get("hasAttachments", False)
                })
            
            return {"emails": emails, "count": len(emails)}
            
    except Exception as e:
        return {"error": str(e)}


async def execute_get_email_content(token: str, message_id: str) -> dict:
    """Get full email content."""
    try:
        async with httpx.AsyncClient() as client:
            print(f"[Graph] get_email_content id={message_id}", file=sys.stderr)
            response = await client.get(
                f"https://graph.microsoft.com/v1.0/me/messages/{message_id}?$select=id,subject,from,toRecipients,receivedDateTime,body",
                headers={"Authorization": f"Bearer {token}"},
                timeout=30.0
            )
            
            if response.status_code != 200:
                print(f"[Graph] get_email_content failed status={response.status_code} body={response.text[:300]}", file=sys.stderr)
                return {"error": f"Failed to get email: {response.status_code}"}
            
            msg = response.json()
            
            # Extract text from HTML body
            body_content = msg.get("body", {}).get("content", "")
            text_body = re.sub(r'<[^>]+>', ' ', body_content)
            text_body = re.sub(r'\s+', ' ', text_body).strip()
            
            return {
                "subject": msg.get("subject", "(No subject)"),
                "from": msg.get("from", {}).get("emailAddress", {}).get("name", "Unknown"),
                "from_email": msg.get("from", {}).get("emailAddress", {}).get("address", ""),
                "date": msg.get("receivedDateTime", ""),
                "body": text_body[:3000]
            }
            
    except Exception as e:
        return {"error": str(e)}


async def execute_search_calendar(token: str, query: str = "", days_ahead: int = 7) -> dict:
    """Search calendar events."""
    try:
        start = datetime.utcnow()
        end = start + timedelta(days=days_ahead)
        
        async with httpx.AsyncClient() as client:
            url = f"https://graph.microsoft.com/v1.0/me/calendarView?startDateTime={start.isoformat()}Z&endDateTime={end.isoformat()}Z&$select=id,subject,start,end,location,organizer,isAllDay&$orderby=start/dateTime&$top=50"
            print(f"[Graph] search_calendar url={url} query={query}", file=sys.stderr)
            
            response = await client.get(
                url,
                headers={"Authorization": f"Bearer {token}"},
                timeout=30.0
            )
            
            if response.status_code != 200:
                print(f"[Graph] search_calendar failed status={response.status_code} body={response.text[:300]}", file=sys.stderr)
                return {"error": f"Failed to search calendar: {response.status_code}"}
            
            data = response.json()
            events = []
            
            for event in data.get("value", []):
                if query:
                    query_lower = query.lower()
                    subject = event.get("subject", "").lower()
                    location = str(event.get("location", {}).get("displayName", "")).lower()
                    if query_lower not in subject and query_lower not in location:
                        continue
                
                events.append({
                    "id": event.get("id"),
                    "subject": event.get("subject", "(No title)"),
                    "start": event.get("start", {}).get("dateTime", ""),
                    "end": event.get("end", {}).get("dateTime", ""),
                    "location": event.get("location", {}).get("displayName", ""),
                    "isAllDay": event.get("isAllDay", False),
                    "organizer": event.get("organizer", {}).get("emailAddress", {}).get("name", "")
                })
            
            return {"events": events, "count": len(events), "days_searched": days_ahead}
            
    except Exception as e:
        return {"error": str(e)}


async def execute_get_todays_events(token: str) -> dict:
    """Get today's calendar events."""
    try:
        now = datetime.utcnow()
        start = now.replace(hour=0, minute=0, second=0, microsecond=0)
        end = start + timedelta(days=1)
        
        async with httpx.AsyncClient() as client:
            url = f"https://graph.microsoft.com/v1.0/me/calendarView?startDateTime={start.isoformat()}Z&endDateTime={end.isoformat()}Z&$select=id,subject,start,end,location,organizer,isAllDay&$orderby=start/dateTime"
            print(f"[Graph] get_todays_events url={url}", file=sys.stderr)
            
            response = await client.get(
                url,
                headers={"Authorization": f"Bearer {token}"},
                timeout=30.0
            )
            
            if response.status_code != 200:
                print(f"[Graph] get_todays_events failed status={response.status_code} body={response.text[:300]}", file=sys.stderr)
                return {"error": f"Failed to get calendar: {response.status_code}"}
            
            data = response.json()
            events = []
            
            for event in data.get("value", []):
                events.append({
                    "subject": event.get("subject", "(No title)"),
                    "start": event.get("start", {}).get("dateTime", ""),
                    "end": event.get("end", {}).get("dateTime", ""),
                    "location": event.get("location", {}).get("displayName", ""),
                    "isAllDay": event.get("isAllDay", False)
                })
            
            if not events:
                return {"message": "No meetings or events scheduled for today.", "events": [], "count": 0}
            
            return {"events": events, "count": len(events), "date": start.strftime("%A, %B %d, %Y")}
            
    except Exception as e:
        return {"error": str(e)}


async def execute_search_files(token: str, query: str) -> dict:
    """Search OneDrive files."""
    try:
        async with httpx.AsyncClient() as client:
            print(f"[Graph] search_files query={query}", file=sys.stderr)
            response = await client.get(
                f"https://graph.microsoft.com/v1.0/me/drive/root/search(q='{query}')?$select=id,name,webUrl,createdDateTime,size,file,folder&$top=20",
                headers={"Authorization": f"Bearer {token}"},
                timeout=30.0
            )
            
            if response.status_code != 200:
                print(f"[Graph] search_files failed status={response.status_code} body={response.text[:300]}", file=sys.stderr)
                return {"error": f"Failed to search files: {response.status_code}"}
            
            data = response.json()
            files = []
            
            for item in data.get("value", []):
                files.append({
                    "id": item.get("id"),
                    "name": item.get("name"),
                    "url": item.get("webUrl"),
                    "created": item.get("createdDateTime"),
                    "size": item.get("size", 0),
                    "isFolder": "folder" in item
                })
            
            return {"files": files, "count": len(files)}
            
    except Exception as e:
        return {"error": str(e)}


async def execute_tool(tool_name: str, arguments: dict, token: str, request: Optional[Request] = None) -> str:
    """Execute a tool and return the result as a string."""
    try:
        print(f"execute_tool start tool={tool_name} args_keys={list(arguments.keys())} token_present={bool(token)}", file=sys.stderr)
        if tool_name == "search_emails":
            result = await execute_search_emails(
                token,
                arguments.get("query", ""),
                arguments.get("limit", 10),
                request=request,
            )
        elif tool_name == "get_email_content":
            result = await execute_get_email_content(
                token,
                arguments.get("message_id", "")
            )
        elif tool_name == "search_calendar":
            result = await execute_search_calendar(
                token,
                arguments.get("query", ""),
                arguments.get("days_ahead", 7)
            )
        elif tool_name == "get_todays_events":
            result = await execute_get_todays_events(token)
        elif tool_name == "search_files":
            result = await execute_search_files(
                token,
                arguments.get("query", "")
            )
        else:
            result = {"error": f"Unknown tool: {tool_name}"}
        
        print(f"execute_tool result tool={tool_name} keys={list(result.keys())}", file=sys.stderr)
        return json.dumps(result, indent=2)
        
    except Exception as e:
        print(f"execute_tool error tool={tool_name}: {e}", file=sys.stderr)
        return json.dumps({"error": str(e)})


# ============== Page Routes ==============

@app.get("/", response_class=HTMLResponse)
async def root(request: Request):
    """Root route - show login or redirect to dashboard."""
    if is_authenticated(request):
        return templates.TemplateResponse("dashboard.html", {"request": request})
    return templates.TemplateResponse("login.html", {"request": request})


@app.get("/login", response_class=HTMLResponse)
async def login_page(request: Request):
    """Login page."""
    if is_authenticated(request):
        return HTMLResponse(
            content='<script>window.location.href="/dashboard";</script>',
            status_code=200
        )
    return templates.TemplateResponse("login.html", {"request": request})


@app.get("/dashboard", response_class=HTMLResponse)
async def dashboard(request: Request):
    """Dashboard - main hub after login."""
    if not is_authenticated(request):
        return HTMLResponse(
            content='<script>window.location.href="/.auth/login/aad?post_login_redirect_uri=/dashboard";</script>',
            status_code=200
        )
    return templates.TemplateResponse("dashboard.html", {"request": request})


@app.get("/chat", response_class=HTMLResponse)
async def chat_page(request: Request):
    """Chat interface with Joogni AI."""
    if not is_authenticated(request):
        return HTMLResponse(
            content='<script>window.location.href="/.auth/login/aad?post_login_redirect_uri=/chat";</script>',
            status_code=200
        )
    deploy_sha = _get_deploy_sha()
    return templates.TemplateResponse("index.html", {
        "request": request,
        "deploy_sha": deploy_sha,
    })


@app.get("/calculators", response_class=HTMLResponse)
async def calculators_page(request: Request):
    """Family law calculators."""
    if not is_authenticated(request):
        return HTMLResponse(
            content='<script>window.location.href="/.auth/login/aad?post_login_redirect_uri=/calculators";</script>',
            status_code=200
        )
    return templates.TemplateResponse("calculators.html", {"request": request})


# ============== API Routes ==============

# In-memory storage for agreements (use database in production)
user_agreements = {}


def _get_deploy_sha() -> str:
    """Read deploy SHA from DEPLOY_SHA.txt for cache busting and verification."""
    try:
        sha_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "DEPLOY_SHA.txt")
        if os.path.isfile(sha_path):
            with open(sha_path) as f:
                return f.read().strip()[:12]
    except Exception:
        pass
    return "not-deployed"


@app.get("/api/version")
async def api_version():
    """Deploy verification: returns version info and deploy SHA. No auth required."""
    deploy_sha = _get_deploy_sha()
    version_file = os.path.join(os.path.dirname(os.path.abspath(__file__)), "static", "version.json")
    if os.path.isfile(version_file):
        try:
            with open(version_file) as f:
                data = json.load(f)
                data["deploy_sha"] = deploy_sha
                return JSONResponse(data)
        except Exception:
            pass
    return JSONResponse({
        "deploy_check": "typography-v1",
        "deploy_sha": deploy_sha,
        "message": "If you see this, app.py was deployed. static/version.json not found.",
    })


@app.get("/api/user")
async def get_user(request: Request):
    """Get current user info."""
    user_info = get_user_info(request)
    if user_info:
        return JSONResponse(user_info)
    return JSONResponse({"authenticated": False})


@app.get("/api/user-info")
async def get_user_info_api(request: Request):
    """Get user info for display."""
    user_info = get_user_info(request)
    if user_info:
        return JSONResponse({
            "name": user_info.get("name", "User"),
            "email": user_info.get("email", ""),
            "authenticated": True
        })
    return JSONResponse({
        "name": "Guest",
        "email": "",
        "authenticated": False
    })


@app.get("/api/check-agreement")
async def check_agreement(request: Request):
    """Check if user has accepted the agreement."""
    user_info = get_user_info(request)
    user_id = user_info.get("email", "anonymous") if user_info else "anonymous"
    
    accepted = user_agreements.get(user_id, False)
    return JSONResponse({"accepted": accepted})


@app.post("/api/accept-agreement")
async def accept_agreement(request: Request):
    """Record user's acceptance of the agreement."""
    user_info = get_user_info(request)
    user_id = user_info.get("email", "anonymous") if user_info else "anonymous"
    
    user_agreements[user_id] = True
    return JSONResponse({"success": True, "message": "Agreement accepted"})


@app.get("/api/box/status")
async def box_status(request: Request):
    """Check Box integration status (placeholder)."""
    return JSONResponse({
        "connected": False,
        "message": "Box integration not configured"
    })

def pick_search_query(messages: List[dict]) -> str:
    """
    Uses the first user question as the search query (the most substantive one).
    For follow-ups like 'add more arguments' or 'make a summary', search uses the original question.
    Fallback: last user message if none is substantive.
    """
    first_substantive = ""
    last_user = ""
    for msg in messages:
        if msg.get("role") == "user":
            text = normalize_to_text(msg.get("content", ""))
            last_user = text
            if len(text.strip()) > 10 and not first_substantive:
                first_substantive = text
    return first_substantive or last_user


def normalize_to_text(x) -> str:
    if x is None:
        return ""
    if isinstance(x, str):
        return x
    if isinstance(x, bytes):
        return x.decode("utf-8", errors="ignore")
    if isinstance(x, list):
        # ex: [{"type":"text","text":"..."}, ...]
        parts = []
        for it in x:
            if isinstance(it, str):
                parts.append(it)
            elif isinstance(it, dict):
                # common patterns
                if "text" in it and isinstance(it["text"], str):
                    parts.append(it["text"])
                elif it.get("type") == "text" and isinstance(it.get("text"), str):
                    parts.append(it["text"])
        return "\n".join([p for p in parts if p])
    if isinstance(x, dict):
        # fallback: try to pick common fields
        if "text" in x and isinstance(x["text"], str):
            return x["text"]
        if "content" in x:
            return normalize_to_text(x["content"])
        return ""
    return str(x)



@app.post("/api/get-upload-url")
async def get_upload_url(request: Request):
    """Return a SAS URL for uploading a single blob (up to MAX_UPLOAD_SIZE_MB)."""
    if not is_authenticated(request):
        raise HTTPException(status_code=401, detail="Not authenticated")

    try:
        data = await request.json()
        filename = (data.get("filename") or "").strip()
        if not filename:
            raise HTTPException(status_code=400, detail="filename is required")

        container = os.getenv("UPLOAD_CONTAINER") or os.getenv("AZURE_STORAGE_CONTAINER") or "legal-docs-raw"
        blob_service = _get_blob_service_client()
        _ensure_container(blob_service, container)

        user_prefix = _get_user_blob_prefix(request)
        blob_name = _safe_blob_name(filename, user_prefix=user_prefix)
        upload_url, final_blob_name = _generate_sas_url(blob_service, container, blob_name)

        return JSONResponse({
            "upload_url": upload_url,
            "blob_name": final_blob_name,
            "max_size_mb": MAX_UPLOAD_SIZE_MB
        })
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


def _user_owns_blob(request: Request, blob_name: str) -> bool:
    """Verify blob belongs to the current user (must have user prefix)."""
    if not blob_name or not isinstance(blob_name, str):
        return False
    prefix = _get_user_blob_prefix(request)
    # Blob must start with "u_xxxx/" to be user-scoped
    return blob_name.startswith(prefix + "/")


@app.post("/api/delete-attachment")
async def delete_attachment(request: Request):
    """Delete blob(s) from storage when user removes PDF or page reloads.
    Only deletes blobs that belong to the current user (user-prefixed paths).
    """
    if not is_authenticated(request):
        raise HTTPException(status_code=401, detail="Not authenticated")

    try:
        try:
            data = await request.json()
        except Exception:
            data = {}
        blob_name = data.get("blob_name")
        blob_names = data.get("blob_names") or []

        if blob_name:
            blob_names = [blob_name]

        if not blob_names:
            return JSONResponse({"deleted": 0, "message": "No blobs to delete"})

        storage = LegalDocsStorage()
        deleted = 0
        for name in blob_names:
            if name and isinstance(name, str) and _user_owns_blob(request, name):
                if storage.delete_file(name):
                    deleted += 1

        return JSONResponse({"deleted": deleted, "blob_names": blob_names})
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/register-attachment")
async def register_attachment(request: Request, background_tasks: BackgroundTasks):
    """Register an uploaded blob for background processing. Returns job_id."""
    if not is_authenticated(request):
        raise HTTPException(status_code=401, detail="Not authenticated")

    print("[register-attachment] endpoint called", file=sys.stderr)
    search_svc = os.getenv("AZURE_SEARCH_SERVICE")
    if search_svc:
        conn_err = _check_search_connectivity(search_svc, label="register-attachment")
        if conn_err:
            print(f"[register-attachment] WARNING: Search unreachable from main request: {conn_err}", file=sys.stderr)
    try:
        data = await request.json()
        blob_name = (data.get("blob_name") or "").strip()
        original_filename = (data.get("original_filename") or "file.pdf").strip()
        if not blob_name:
            raise HTTPException(status_code=400, detail="blob_name is required")
        if not _user_owns_blob(request, blob_name):
            raise HTTPException(status_code=403, detail="Blob does not belong to user")

        job_id = str(uuid.uuid4())
        user_prefix = _get_user_blob_prefix(request)
        if not create_job(user_prefix, job_id, blob_name, original_filename):
            raise HTTPException(status_code=500, detail="Failed to create job")

        background_tasks.add_task(process_attachment_job, job_id, user_prefix, blob_name, original_filename)
        return JSONResponse({
            "job_id": job_id,
            "blob_name": blob_name,
            "status": "pending",
        })
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/job-storage-status")
async def job_storage_status():
    """Return which backend is used for job storage: azure_table or in_memory."""
    backend = get_job_storage_backend()
    return JSONResponse({"backend": backend, "using_azure_table": backend == "azure_table"})


@app.get("/api/jobs/{job_id}")
async def get_job_status(job_id: str, request: Request):
    """Get status of an attachment processing job."""
    if not is_authenticated(request):
        raise HTTPException(status_code=401, detail="Not authenticated")
    user_prefix = _get_user_blob_prefix(request)
    job = get_job(user_prefix, job_id)
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")
    status = job.get("status", "unknown")
    return JSONResponse({
        "job_id": job_id,
        "status": status,
        "blob_name": job.get("blob_name"),
        "original_filename": job.get("original_filename"),
        "error": job.get("error") if status == "failed" else None,
    })


async def _run_agentic_conversation(
    job_id: str, user_prefix: str, data: dict, graph_token: Optional[str]
) -> None:
    """Background task: run the full agentic conversation and store result in job."""
    import sys
    partition = "agt_" + user_prefix
    try:
        update_agentic_job(partition, job_id, "processing")
        response_text = await _execute_agentic_logic(data, user_prefix, graph_token)
        update_agentic_job(partition, job_id, "completed", response=response_text)
    except HTTPException as e:
        update_agentic_job(partition, job_id, "failed", error=str(e.detail))
    except Exception as e:
        import traceback
        traceback.print_exc(file=sys.stderr)
        update_agentic_job(partition, job_id, "failed", error=str(e))


def _stream_llm_to_queue_sync(client, model: str, messages: list, queue, loop, temperature: float = 0.7):
    """Run LLM stream in sync context; put chunks in asyncio.Queue from a thread."""
    stream = client.chat.completions.create(
        model=model,
        messages=messages,
        temperature=temperature,
        max_tokens=2000,
        stream=True,
    )
    for chunk in stream:
        content = (chunk.choices[0].delta.content or "") if chunk.choices else ""
        if content:
            loop.call_soon_threadsafe(queue.put_nowait, content)
    loop.call_soon_threadsafe(queue.put_nowait, None)


async def _stream_llm_to_queue_async(async_client: AsyncAzureOpenAI, model: str, messages: list, queue: asyncio.Queue, temperature: float = 0.7):
    """Stream LLM tokens into queue using async client — no threads, true async concurrency."""
    stream = await async_client.chat.completions.create(
        model=model,
        messages=messages,
        temperature=temperature,
        max_tokens=2000,
        stream=True,
    )
    async for chunk in stream:
        content = (chunk.choices[0].delta.content or "") if chunk.choices else ""
        if content:
            await queue.put(content)
    await queue.put(None)


async def _stream_with_tools_to_queue_async(
    async_client: AsyncAzureOpenAI, model: str, messages: list,
    queue: asyncio.Queue, tools: list, graph_token: Optional[str], temperature: float = 0.7
):
    """Stream first LLM call with tools. Forwards content chunks immediately.
    If the model invokes tools, executes them and streams the 2nd response.
    """
    stream = await async_client.chat.completions.create(
        model=model,
        messages=messages,
        tools=tools,
        tool_choice="auto",
        temperature=temperature,
        max_tokens=2000,
        stream=True,
    )

    tool_calls_acc: dict = {}  # index -> {id, name, arguments}
    assistant_content = ""

    async for chunk in stream:
        if not chunk.choices:
            continue
        delta = chunk.choices[0].delta

        if delta.content:
            await queue.put(delta.content)
            assistant_content += delta.content

        if delta.tool_calls:
            for tc in delta.tool_calls:
                idx = tc.index
                if idx not in tool_calls_acc:
                    tool_calls_acc[idx] = {"id": "", "name": "", "arguments": ""}
                if tc.id:
                    tool_calls_acc[idx]["id"] = tc.id
                if tc.function:
                    if tc.function.name:
                        tool_calls_acc[idx]["name"] += tc.function.name
                    if tc.function.arguments:
                        tool_calls_acc[idx]["arguments"] += tc.function.arguments

    # No tool calls — content was already streamed
    if not tool_calls_acc or not graph_token:
        await queue.put(None)
        return

    # Build tool call list and append assistant turn
    tool_calls_list = [
        {"id": v["id"], "type": "function", "function": {"name": v["name"], "arguments": v["arguments"]}}
        for v in tool_calls_acc.values()
    ]
    messages.append({
        "role": "assistant",
        "content": assistant_content or None,
        "tool_calls": tool_calls_list,
    })

    for tc in tool_calls_list:
        tool_result = await execute_tool(tc["function"]["name"], json.loads(tc["function"]["arguments"]), graph_token, None)
        messages.append({"role": "tool", "tool_call_id": tc["id"], "content": tool_result})

    # Stream 2nd response after tool execution
    await _stream_llm_to_queue_async(async_client, model, messages, queue, temperature=0.7)


async def _execute_agentic_logic(
    data: dict, user_prefix: str, graph_token: Optional[str],
    stream_queue: Optional[asyncio.Queue] = None,
) -> str:
    """Execute agentic conversation logic; returns response text or raises."""
    import sys
    jurisdiction_raw = data.get("jurisdiction", "California")
    jid, j_config = resolve_jurisdiction(jurisdiction_raw)
    messages = data.get("messages", [])
    if not messages:
        messages = data.get("conversation", [])
    if not messages:
        messages = data.get("history", [])
    if not messages and "message" in data:
        messages = [{"role": "user", "content": data.get("message", "")}]
    if not messages and "query" in data:
        messages = [{"role": "user", "content": data.get("query", "")}]
    if not messages and "prompt" in data:
        messages = [{"role": "user", "content": data.get("prompt", "")}]

    context = data.get("context", "")
    print(f"Messages after parsing: {len(messages)}", file=sys.stderr)

    key = os.getenv("AZURE_OPENAI_KEY")
    raw_endpoint = os.getenv("AZURE_OPENAI_ENDPOINT")
    deployment = os.getenv("AZURE_OPENAI_DEPLOYMENT", "gpt-4.1")
    search_service = os.getenv("AZURE_SEARCH_SERVICE")
    search_index = os.getenv("AZURE_SEARCH_INDEX")
    search_key = os.getenv("AZURE_SEARCH_KEY")
    endpoint = raw_endpoint
    version = os.getenv("AZURE_OPENAI_API_VERSION", "2023-05-15")

    if not key or not endpoint:
        raise HTTPException(status_code=500, detail="Azure OpenAI not configured")

    print("Creating AzureOpenAI client...", file=sys.stderr)
    client = AzureOpenAI(
        api_key=key,
        azure_endpoint=endpoint,
        api_version=version,
    )
    async_client = AsyncAzureOpenAI(
        api_key=key,
        azure_endpoint=endpoint,
        api_version=version,
    )
    print("Client created, making request...", file=sys.stderr)

    free_mode = data.get("free_mode")
    if free_mode:
        system_prompt = FREE_CHAT_SYSTEM_PROMPT
        chat_messages = [{"role": "system", "content": system_prompt}]
        for msg in messages:
            chat_messages.append({
                "role": msg.get("role", "user"),
                "content": msg.get("content", "")
            })
        tools_for_request = None
        temperature = 0.5
        legal_mode = False
        print("Free mode: using generic prompt, no RAG, no tools", file=sys.stderr)
    else:
        system_prompt = JURISDICTION_SYSTEM_PROMPT.get(jid) or ""
        chat_messages = [{"role": "system", "content": system_prompt}]

        search_query = pick_search_query(messages)
        has_prior_context = any(m.get("role") == "assistant" for m in messages) or bool(data.get("context"))

        blobs = data.get("blobs", [])
        print(f"DEBUG: blobs count={len(blobs)} keys={list(data.keys())}", file=sys.stderr)
        for b in blobs:
            fn = (b.get("original_filename") or b.get("blob_name") or "")
            if not _is_supported_attachment(fn):
                raise HTTPException(
                    status_code=400,
                    detail="Unsupported file type. Only PDF and Word (.docx) are allowed.",
                )

        legal_mode = False
        context_chunks = []
        no_context = False
        attachment_contexts: List[str] = []
        search_client = None
        if search_service and search_index and search_key:
            search_endpoint = f"https://{search_service}.search.windows.net"
            search_client = SearchClient(
                search_endpoint, search_index, AzureKeyCredential(search_key)
            )
        if search_client and search_query:
            try:
                request_domains = data.get("domains", [])
                if isinstance(request_domains, str):
                    request_domains = [d.strip() for d in request_domains.split(",") if d.strip()]
                elif not isinstance(request_domains, list):
                    request_domains = []

                candidates = request_domains or j_config.get("domains") or LEGAL_DOMAINS
                candidates = [d for d in candidates if d][:LEGAL_MAX_DOMAINS]
                search_query = normalize_to_text(search_query)

                t0 = time.perf_counter()
                print(f"Running hybrid search for query='{search_query[:80]}...' domains={candidates}", file=sys.stderr)
                context_chunks, domains_with_hits = await run_multi_domain_search_async(
                    search_client=search_client,
                    query=search_query,
                    domains=candidates,
                    max_domains=LEGAL_MAX_DOMAINS,
                    top_per_domain=LEGAL_TOP_PER_DOMAIN,
                    embed_client=client,
                )
                print(f"[perf] run_multi_domain_search: {time.perf_counter() - t0:.2f}s", file=sys.stderr)

                legal_mode = len(context_chunks) > 0
                print(f"Legal mode via search hits: {legal_mode} (domains_with_hits={domains_with_hits} chunks={len(context_chunks)})", file=sys.stderr)
                if legal_mode:
                    ctx = "\n\n".join(context_chunks)
                    pref = "When the user has attachments, treat them as the primary source; use the context below only if relevant to the attachment's topic.\n\n" if blobs else ""
                    chat_messages.append({
                        "role": "user",
                        "content": (
                            pref
                            + "Use the context below (tagged by DOMAIN). Cite the sources. "
                            "Reuse and expand upon prior conversation content when the user asks for more details, expansion, or revision.\n"
                            + ctx
                        )
                    })
                    print(f"Legal RAG applied with {len(context_chunks)} chunk(s)", file=sys.stderr)
                else:
                    no_context = True
                    print("Legal RAG skipped (no relevant search hits)", file=sys.stderr)
            except Exception as search_err:
                print(f"Legal RAG search error: {search_err}", file=sys.stderr)
        else:
            print("Legal RAG not applied (missing search config or empty question)", file=sys.stderr)

        if blobs:
            blobs_with_job = [b for b in blobs if b.get("job_id")]
            blobs_without_job = [b for b in blobs if not b.get("job_id")]
            if blobs_without_job:
                names = [b.get("original_filename", "file") for b in blobs_without_job]
                raise HTTPException(
                    status_code=400,
                    detail=f"File(s) '{', '.join(names)}' were not processed. Remove and re-upload to process in the background.",
                )
            processing_names = []
            failed_jobs = []
            full_text_blobs: List[dict] = []
            rag_blob_names: List[str] = []
            for b in blobs_with_job:
                att_job_id = b.get("job_id")
                blob_name = b.get("blob_name")
                orig = b.get("original_filename", "file.pdf")
                if not att_job_id or not blob_name:
                    continue
                job = get_job(user_prefix, att_job_id)
                if not job:
                    processing_names.append(orig)
                    continue
                st = job.get("status", "pending")
                if st == "completed":
                    if job.get("use_full_text") == "true":
                        full_text_blobs.append(b)
                    else:
                        rag_blob_names.append(blob_name)
                elif st == "failed":
                    failed_jobs.append((orig, job.get("error", "Unknown error")))
                else:
                    processing_names.append(orig)
            if processing_names:
                msg = f"Document(s) still processing: {', '.join(processing_names)}. Please wait a few seconds and try again."
                raise HTTPException(status_code=400, detail=msg)
            if failed_jobs:
                err_parts = [f"{n}: {e}" for n, e in failed_jobs]
                raise HTTPException(status_code=400, detail=f"Attachment processing failed: {'; '.join(err_parts)}")
            if full_text_blobs:
                try:
                    t0_ft = time.perf_counter()
                    ft_sections = await _fetch_full_text_context(full_text_blobs)
                    if ft_sections:
                        attachment_contexts.extend(ft_sections)
                    print(f"[perf] _fetch_full_text_context: {time.perf_counter() - t0_ft:.2f}s for {len(full_text_blobs)} blob(s)", file=sys.stderr)
                except Exception as ft_err:
                    print(f"[UserAttachment] Full-text fetch error: {ft_err}", file=sys.stderr)
            if rag_blob_names and search_client:
                try:
                    t0_attach = time.perf_counter()
                    attach_chunks = await run_user_attachment_search_async(
                        search_client=search_client,
                        query=search_query or "legal",
                        blob_names=rag_blob_names,
                        top_k=6,
                        embed_client=client,
                    )
                    print(f"[perf] run_user_attachment_search: {time.perf_counter() - t0_attach:.2f}s", file=sys.stderr)
                    if attach_chunks:
                        attachment_contexts.extend(attach_chunks)
                    print(f"[UserAttachment] RAG hit {len(attach_chunks)} chunks for {len(rag_blob_names)} blob(s)", file=sys.stderr)
                except Exception as ua_err:
                    print(f"[UserAttachment] Search error: {ua_err}", file=sys.stderr)

        if attachment_contexts:
            no_context = False
            chat_messages.append({
                "role": "user",
                "content": (
                    "Use the following attachment excerpts to answer. "
                    "Provide a separate section for EACH attachment, even if you just say it is not relevant. "
                    "Cite them as 'Attachment Source'.\n\n"
                    + "\n\n".join(attachment_contexts)
                )
            })

        if has_prior_context:
            no_context = False

        if no_context:
            chat_messages.append({
                "role": "user",
                "content": (
                    "No relevant matches were found in the legal index for this specific query. "
                    "You may answer based on prior conversation context, attachments, or indicate that no matching statutes/cases were found. "
                    "Never refuse with 'I don't have enough information' if you have context from the conversation or attachments."
                )
            })

        if context:
            chat_messages.append({
                "role": "user",
                "content": f"Context from uploaded documents/emails:\n{context}"
            })
            chat_messages.append({
                "role": "assistant",
                "content": "I've reviewed the provided context. How can I help you with this information?"
            })

        for msg in messages:
            chat_messages.append({
                "role": msg.get("role", "user"),
                "content": msg.get("content", "")
            })

        print(f"Total chat_messages: {len(chat_messages)}", file=sys.stderr)
        temperature = 0.2 if (legal_mode and not has_prior_context) else 0.5
        tools_for_request = get_tools(jid, j_config, bool(graph_token))
        print(f"Making Azure OpenAI call with tools={bool(tools_for_request)} legal_mode={legal_mode} temp={temperature}...", file=sys.stderr)

    try:
        t0_llm = time.perf_counter()
        if stream_queue:
            if tools_for_request:
                await _stream_with_tools_to_queue_async(async_client, deployment, chat_messages, stream_queue, tools_for_request, graph_token, temperature)
                print(f"[perf] chat.completions.create (streamed, with tools): {time.perf_counter() - t0_llm:.2f}s", file=sys.stderr)
            else:
                await _stream_llm_to_queue_async(async_client, deployment, chat_messages, stream_queue, temperature)
                print(f"[perf] chat.completions.create (streamed): {time.perf_counter() - t0_llm:.2f}s", file=sys.stderr)
            return ""
        response = await async_client.chat.completions.create(
            model=deployment,
            messages=chat_messages,
            tools=tools_for_request,
            tool_choice="auto" if tools_for_request else None,
            temperature=temperature,
            max_tokens=2000,
        )
        print(f"[perf] chat.completions.create (1st): {time.perf_counter() - t0_llm:.2f}s", file=sys.stderr)
    except Exception as api_error:
        print(f"Azure OpenAI API Error: {type(api_error).__name__}: {api_error}", file=sys.stderr)
        raise

    response_message = response.choices[0].message

    if response_message.tool_calls and graph_token:
        chat_messages.append({
            "role": "assistant",
            "content": response_message.content,
            "tool_calls": [
                {
                    "id": tc.id,
                    "type": "function",
                    "function": {
                        "name": tc.function.name,
                        "arguments": tc.function.arguments
                    }
                }
                for tc in response_message.tool_calls
            ]
        })

        for tool_call in response_message.tool_calls:
            function_name = tool_call.function.name
            function_args = json.loads(tool_call.function.arguments)
            tool_result = await execute_tool(function_name, function_args, graph_token, None)
            chat_messages.append({
                "role": "tool",
                "tool_call_id": tool_call.id,
                "content": tool_result
            })

        t0_llm2 = time.perf_counter()
        final_response = await async_client.chat.completions.create(
            model=deployment,
            messages=chat_messages,
            temperature=0.7,
            max_tokens=2000,
        )
        print(f"[perf] chat.completions.create (2nd, after tools): {time.perf_counter() - t0_llm2:.2f}s", file=sys.stderr)
        return final_response.choices[0].message.content or ""

    return response_message.content or ""


@app.post("/api/conversation")
@app.post("/api/agentic")
async def conversation(request: Request, background_tasks: BackgroundTasks):
    """Handle chat - returns job_id immediately; processing runs in background."""
    import sys
    print("=== CONVERSATION ENDPOINT HIT ===", file=sys.stderr)
    try:
        data = await request.json()
        jurisdiction_raw = data.get("jurisdiction", "California")
        jid, j_config = resolve_jurisdiction(jurisdiction_raw)
        display_name = j_config.get("display_name", jurisdiction_raw)
        if not j_config.get("implemented"):
            return JSONResponse({
                "job_id": str(uuid.uuid4()),
                "status": "completed",
                "response": f"This state ({display_name}) is not yet supported. Only California is currently available.",
                "error": "STATE_NOT_SUPPORTED"
            })
        user_prefix = _get_user_blob_prefix(request)
        graph_token = await get_valid_graph_token(request)
        blobs = data.get("blobs", [])
        if blobs:
            for b in blobs:
                fn = (b.get("original_filename") or b.get("blob_name") or "")
                if not _is_supported_attachment(fn):
                    raise HTTPException(status_code=400, detail="Unsupported file type. Only PDF and Word (.docx) are allowed.")
            blobs_with_job = [b for b in blobs if b.get("job_id")]
            blobs_without_job = [b for b in blobs if not b.get("job_id")]
            if blobs_without_job:
                names = [b.get("original_filename", "file") for b in blobs_without_job]
                raise HTTPException(
                    status_code=400,
                    detail=f"File(s) '{', '.join(names)}' were not processed. Remove and re-upload to process in the background.",
                )
            search_query = pick_search_query(data.get("messages", []) or [{"content": data.get("query", "")}])
            processing_names = []
            failed_jobs = []
            for b in blobs_with_job:
                att_job_id = b.get("job_id")
                orig = b.get("original_filename", "file.pdf")
                if not att_job_id:
                    continue
                job = get_job(user_prefix, att_job_id)
                if not job:
                    processing_names.append(orig)
                elif job.get("status") == "failed":
                    failed_jobs.append((orig, job.get("error", "Unknown error")))
                elif job.get("status") != "completed":
                    processing_names.append(orig)
            if processing_names:
                msg = f"Document(s) still processing: {', '.join(processing_names)}. Please wait and try again."
                return JSONResponse({"job_id": str(uuid.uuid4()), "status": "processing", "response": msg})
            if failed_jobs:
                err_parts = [f"{n}: {e}" for n, e in failed_jobs]
                raise HTTPException(status_code=400, detail=f"Attachment processing failed: {'; '.join(err_parts)}")
        key = os.getenv("AZURE_OPENAI_KEY")
        endpoint = os.getenv("AZURE_OPENAI_ENDPOINT")
        if not key or not endpoint:
            raise HTTPException(status_code=500, detail="Azure OpenAI not configured")
        job_id = str(uuid.uuid4())
        partition = "agt_" + user_prefix
        if not create_agentic_job(partition, job_id):
            raise HTTPException(status_code=500, detail="Failed to create job")
        background_tasks.add_task(_run_agentic_conversation, job_id, user_prefix, data, graph_token)
        return JSONResponse({"job_id": job_id, "status": "pending"})
    except HTTPException:
        raise
    except Exception as e:
        import traceback
        traceback.print_exc(file=sys.stderr)
        raise HTTPException(status_code=500, detail=str(e))


async def _stream_agentic_events(
    data: dict, user_prefix: str, graph_token: Optional[str]
) -> AsyncGenerator[str, None]:
    """Yield SSE-formatted events. Sends done event with complete=true for truncation detection."""
    queue: asyncio.Queue = asyncio.Queue()
    task = asyncio.create_task(
        _execute_agentic_logic(data, user_prefix, graph_token, stream_queue=queue)
    )
    # Flush Azure App Service nginx proxy buffer so tokens stream in real-time.
    # nginx default proxy_buffers total ~64KB; send 80KB to guarantee overflow.
    for _ in range(20):
        yield f":{' ' * 4096}\n\n"
    try:
        while True:
            try:
                chunk = await asyncio.wait_for(queue.get(), timeout=30.0)
            except asyncio.TimeoutError:
                # Heartbeat prevents Gateway Timeout during long LLM calls
                yield ": keep-alive\n\n"
                continue
            if chunk is None:
                break
            yield f"data: {json.dumps({'type': 'token', 'content': chunk})}\n\n"
        yield f"data: {json.dumps({'type': 'done', 'complete': True})}\n\n"
    except asyncio.CancelledError:
        task.cancel()
        try:
            await task
        except asyncio.CancelledError:
            pass
        raise
    finally:
        await task


@app.post("/api/agentic/stream")
async def conversation_stream(request: Request):
    """
    Streaming chat via SSE. Same input as /api/agentic but streams tokens.
    Frontend should detect 'done' event with complete=true; if absent, response may be truncated.
    """
    try:
        data = await request.json()
        jurisdiction_raw = data.get("jurisdiction", "California")
        jid, j_config = resolve_jurisdiction(jurisdiction_raw)
        display_name = j_config.get("display_name", jurisdiction_raw)
        if not j_config.get("implemented"):
            return JSONResponse({
                "error": f"This state ({display_name}) is not yet supported.",
                "code": "STATE_NOT_SUPPORTED"
            }, status_code=400)
        user_prefix = _get_user_blob_prefix(request)
        graph_token = await get_valid_graph_token(request)
        blobs = data.get("blobs", [])
        if blobs:
            for b in blobs:
                fn = (b.get("original_filename") or b.get("blob_name") or "")
                if not _is_supported_attachment(fn):
                    raise HTTPException(status_code=400, detail="Unsupported file type. Only PDF and Word (.docx) are allowed.")
            blobs_with_job = [b for b in blobs if b.get("job_id")]
            blobs_without_job = [b for b in blobs if not b.get("job_id")]
            if blobs_without_job:
                names = [b.get("original_filename", "file") for b in blobs_without_job]
                raise HTTPException(
                    status_code=400,
                    detail=f"File(s) '{', '.join(names)}' were not processed. Remove and re-upload to process in the background.",
                )
            processing_names = []
            for b in blobs_with_job:
                job = get_job(user_prefix, b.get("job_id"))
                if not job or job.get("status") != "completed":
                    processing_names.append(b.get("original_filename", "file"))
            if processing_names:
                raise HTTPException(
                    status_code=400,
                    detail=f"Document(s) still processing: {', '.join(processing_names)}. Please wait and try again.",
                )
        return StreamingResponse(
            _stream_agentic_events(data, user_prefix, graph_token),
            media_type="text/event-stream",
            headers={
                "Cache-Control": "no-cache",
                "Connection": "keep-alive",
                "X-Accel-Buffering": "no",
            },
        )
    except HTTPException:
        raise
    except Exception as e:
        import traceback
        traceback.print_exc(file=sys.stderr)
        raise HTTPException(status_code=500, detail=str(e))


@app.websocket("/api/agentic/ws")
async def agentic_websocket(websocket: WebSocket):
    """WebSocket streaming chat — bypasses Easy Auth proxy buffering."""
    await websocket.accept()
    try:
        data = await websocket.receive_json()
        jurisdiction_raw = data.get("jurisdiction", "California")
        jid, j_config = resolve_jurisdiction(jurisdiction_raw)
        display_name = j_config.get("display_name", jurisdiction_raw)

        if not j_config.get("implemented"):
            await websocket.send_json({"type": "error", "message": f"This state ({display_name}) is not yet supported."})
            return

        user_prefix = _get_user_blob_prefix(websocket)
        graph_token = await get_valid_graph_token(websocket)

        blobs = data.get("blobs", [])
        if blobs:
            for b in blobs:
                fn = (b.get("original_filename") or b.get("blob_name") or "")
                if not _is_supported_attachment(fn):
                    await websocket.send_json({"type": "error", "message": "Unsupported file type. Only PDF and Word (.docx) are allowed."})
                    return
            blobs_without_job = [b for b in blobs if not b.get("job_id")]
            if blobs_without_job:
                names = [b.get("original_filename", "file") for b in blobs_without_job]
                await websocket.send_json({"type": "error", "message": f"File(s) '{', '.join(names)}' were not processed. Remove and re-upload."})
                return
            processing_names = []
            for b in [b for b in blobs if b.get("job_id")]:
                job = get_job(user_prefix, b.get("job_id"))
                if not job or job.get("status") != "completed":
                    processing_names.append(b.get("original_filename", "file"))
            if processing_names:
                await websocket.send_json({"type": "error", "message": f"Document(s) still processing: {', '.join(processing_names)}. Please wait and try again."})
                return

        queue: asyncio.Queue = asyncio.Queue()
        task = asyncio.create_task(
            _execute_agentic_logic(data, user_prefix, graph_token, stream_queue=queue)
        )
        # Guarantee the queue reader always unblocks when the task finishes,
        # even if _execute_agentic_logic returns early without streaming.
        task.add_done_callback(lambda _: queue.put_nowait(None))
        try:
            while True:
                try:
                    chunk = await asyncio.wait_for(queue.get(), timeout=30.0)
                except asyncio.TimeoutError:
                    await websocket.send_json({"type": "ping"})
                    continue
                if chunk is None:
                    break
                await websocket.send_json({"type": "token", "content": chunk})
            await websocket.send_json({"type": "done", "complete": True})
        except WebSocketDisconnect:
            task.cancel()
        finally:
            await task
    except WebSocketDisconnect:
        pass
    except Exception as e:
        import traceback
        traceback.print_exc(file=sys.stderr)
        try:
            await websocket.send_json({"type": "error", "message": str(e)})
        except Exception:
            pass


@app.get("/api/graph/health")
async def graph_health(request: Request):
    """
    Lightweight Graph diagnostic: checks /me and one messages fetch using the token from Easy Auth.
    """
    token = await get_valid_graph_token(request)
    if not token:
        return JSONResponse({"error": "Missing Graph token"}, status_code=401)
    
    results = {}
    async with httpx.AsyncClient(timeout=10.0) as client:
        for name, url in [
            ("me", "https://graph.microsoft.com/v1.0/me"),
            ("messages", "https://graph.microsoft.com/v1.0/me/messages?$top=1&$select=id,subject"),
        ]:
            try:
                resp = await client.get(url, headers={"Authorization": f"Bearer {token}"})
                try:
                    payload = resp.json()
                except Exception:
                    payload = resp.text[:500]
                results[name] = {
                    "status": resp.status_code,
                    "ok": resp.status_code < 300,
                    "payload_preview": payload if resp.status_code >= 300 else None,
                }
                print(f"[Graph health] {name} status={resp.status_code}", file=sys.stderr)
            except Exception as err:
                results[name] = {"status": None, "ok": False, "error": str(err)}
                print(f"[Graph health] {name} error: {err}", file=sys.stderr)
    
    return JSONResponse({"results": results})


@app.get("/api/check_status/{request_id}")
async def check_status(request_id: str, request: Request):
    """Check status of agentic job - returns status, response when completed, or result (error) when failed."""
    user_prefix = _get_user_blob_prefix(request)
    partition = "agt_" + user_prefix
    job = get_job(partition, request_id)
    if not job:
        return JSONResponse({"status": "pending", "request_id": request_id})
    status = (job.get("status") or "pending").lower()
    if status == "failed":
        return JSONResponse({
            "status": "Failed",
            "request_id": request_id,
            "result": job.get("error", "Unknown error"),
        })
    if status == "completed":
        return JSONResponse({
            "status": "completed",
            "request_id": request_id,
            "response": job.get("response", ""),
        })
    return JSONResponse({
        "status": status if status != "pending" else "pending",
        "request_id": request_id,
    })


@app.post("/api/documents/analyze")
async def analyze_document(request: Request, file: UploadFile = File(...)):
    """Analyze uploaded document using Azure Document Intelligence."""
    try:
        from azure.ai.documentintelligence import DocumentIntelligenceClient
        from azure.core.credentials import AzureKeyCredential
        
        doc_endpoint = os.getenv("AZURE_DOCUMENT_INTELLIGENCE_ENDPOINT")
        doc_key = os.getenv("AZURE_DOCUMENT_INTELLIGENCE_KEY")
        
        if not doc_endpoint or not doc_key:
            raise HTTPException(status_code=500, detail="Document Intelligence not configured")
        
        client = DocumentIntelligenceClient(
            endpoint=doc_endpoint,
            credential=AzureKeyCredential(doc_key)
        )
        
        content = await file.read()
        
        poller = client.begin_analyze_document(
            "prebuilt-layout",
            body=content,
            content_type=file.content_type
        )
        result = poller.result()
        
        text_content = ""
        for page in result.pages:
            for line in page.lines:
                text_content += line.content + "\n"
        
        return JSONResponse({
            "filename": file.filename,
            "content": text_content,
            "pages": len(result.pages)
        })
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# ============== Direct Microsoft Graph API Routes (for manual panel) ==============

@app.get("/api/outlook/search")
async def search_emails_direct(request: Request, q: str = "", top: int = 20):
    """Search Outlook emails (direct API for M365 panel)."""
    try:
        token = await get_valid_graph_token(request)
        if not token:
            raise HTTPException(status_code=401, detail="No access token available")
        
        async with httpx.AsyncClient() as client:
            if q:
                url = f"https://graph.microsoft.com/v1.0/me/messages?$search=\"{q}\"&$top={top}&$select=id,subject,from,receivedDateTime,bodyPreview,hasAttachments"
            else:
                url = f"https://graph.microsoft.com/v1.0/me/messages?$top={top}&$orderby=receivedDateTime desc&$select=id,subject,from,receivedDateTime,bodyPreview,hasAttachments"
            
            response = await client.get(
                url,
                headers={"Authorization": f"Bearer {token}"}
            )
            
            if response.status_code != 200:
                raise HTTPException(status_code=response.status_code, detail=response.text)
            
            return JSONResponse(response.json())
            
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/outlook/message/{message_id}")
async def get_email_direct(request: Request, message_id: str):
    """Get full email content (direct API for M365 panel)."""
    try:
        token = await get_valid_graph_token(request)
        if not token:
            raise HTTPException(status_code=401, detail="No access token available")
        
        async with httpx.AsyncClient() as client:
            response = await client.get(
                f"https://graph.microsoft.com/v1.0/me/messages/{message_id}?$select=id,subject,from,toRecipients,receivedDateTime,body,hasAttachments",
                headers={"Authorization": f"Bearer {token}"}
            )
            
            if response.status_code != 200:
                raise HTTPException(status_code=response.status_code, detail=response.text)
            
            return JSONResponse(response.json())
            
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/calendar/search")
async def search_calendar_direct(request: Request, q: str = "", days: int = 30):
    """Search calendar events (direct API for M365 panel)."""
    try:
        token = await get_valid_graph_token(request)
        if not token:
            raise HTTPException(status_code=401, detail="No access token available")
        
        start = datetime.utcnow()
        end = start + timedelta(days=days)
        
        async with httpx.AsyncClient() as client:
            url = f"https://graph.microsoft.com/v1.0/me/calendarView?startDateTime={start.isoformat()}Z&endDateTime={end.isoformat()}Z&$select=id,subject,start,end,location,organizer&$orderby=start/dateTime"
            
            response = await client.get(
                url,
                headers={"Authorization": f"Bearer {token}"}
            )
            
            if response.status_code != 200:
                raise HTTPException(status_code=response.status_code, detail=response.text)
            
            data = response.json()
            
            if q:
                q_lower = q.lower()
                data["value"] = [
                    event for event in data.get("value", [])
                    if q_lower in event.get("subject", "").lower()
                    or q_lower in str(event.get("location", {}).get("displayName", "")).lower()
                ]
            
            return JSONResponse(data)
            
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/onedrive/search")
async def search_onedrive_direct(request: Request, q: str):
    """Search OneDrive files (direct API for M365 panel)."""
    try:
        token = await get_valid_graph_token(request)
        if not token:
            raise HTTPException(status_code=401, detail="No access token available")
        
        async with httpx.AsyncClient() as client:
            response = await client.get(
                f"https://graph.microsoft.com/v1.0/me/drive/root/search(q='{q}')?$select=id,name,webUrl,createdDateTime,size,file",
                headers={"Authorization": f"Bearer {token}"}
            )
            
            if response.status_code != 200:
                raise HTTPException(status_code=response.status_code, detail=response.text)
            
            return JSONResponse(response.json())
            
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# ============== Health Check ==============

@app.get("/health")
async def health_check():
    """Health check endpoint."""
    return {"status": "healthy", "timestamp": datetime.utcnow().isoformat()}


@app.get("/api/test")
async def test_endpoint():
    """Test endpoint to verify API is working."""
    import sys
    print("=== TEST ENDPOINT HIT ===", file=sys.stderr)
    key = os.getenv("AZURE_OPENAI_KEY")
    endpoint = os.getenv("AZURE_OPENAI_ENDPOINT")
    deployment = os.getenv("AZURE_OPENAI_DEPLOYMENT")
    return {
        "status": "ok",
        "key_present": bool(key),
        "endpoint": endpoint,
        "deployment": deployment
    }


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
